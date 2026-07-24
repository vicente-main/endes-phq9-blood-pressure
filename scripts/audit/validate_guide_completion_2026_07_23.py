"""Validación reproducible del cierre de la guía corregida del 23-07-2026.

El informe separa fallos técnicos de la marca ética ``EDITORIAL HOLD``. Esta
última solo desaparece cuando se incorpora al manuscrito una determinación
institucional verificable; no puede resolverse por inferencia del script.
"""

from __future__ import annotations

import hashlib
import json
import math
import re
import sys
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
SUBMISSION = ROOT / "ENVIO_PLOS_ONE_2019-2025"
REVIEWERS = SUBMISSION / "Revisores"
MANUSCRIPT = REVIEWERS / "MANUSCRITO_EN_PLOS_ONE.docx"
MANUSCRIPT_BACKUP = REVIEWERS / "MANUSCRITO_EN_PLOS_ONE.pre_guia_2026-07-23.docx"
COVER = SUBMISSION / "COVER_LETTER.docx"
MODELS = ROOT / "data" / "output_2025" / "analysis" / "models"
TABLES = ROOT / "data" / "output_2025" / "analysis" / "tables"
REPORT_JSON = REVIEWERS / "VALIDACION_GUIA_2026-07-23.json"
REPORT_MD = REVIEWERS / "VALIDACION_GUIA_2026-07-23.md"
ZENODO_ARCHIVE = REVIEWERS / "ZENODO_RELEASE_CANDIDATE_v1.2.0.zip"
ZENODO_CHECKSUM = REVIEWERS / "ZENODO_RELEASE_CANDIDATE_v1.2.0.sha256"
EXPECTED_SOURCE_SHA = "C56E7016BD881D3402529933D2E4A28BC61529A4D9EA4BC2238415D02C926DD7"

DATA_STATEMENT = (
    "The analysis code supporting this study is publicly available on Zenodo at "
    "https://doi.org/10.5281/zenodo.21328300. The de-identified individual-level ENDES "
    "microdata are publicly available without restriction from Peru's National Institute of "
    "Statistics and Informatics (INEI) Microdata Portal at "
    "https://proyectos.inei.gob.pe/microdatos/."
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def workbook_text(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    values: list[str] = []
    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows(values_only=True):
            values.extend(str(value) for value in row if value not in (None, ""))
    return "\n".join(values)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def holm(values: pd.Series) -> pd.Series:
    clean = values.dropna().sort_values()
    total = len(clean)
    result: dict[object, float] = {}
    running = 0.0
    for rank, (index, value) in enumerate(clean.items()):
        adjusted = min(1.0, (total - rank) * float(value))
        running = max(running, adjusted)
        result[index] = running
    return pd.Series(result).reindex(values.index)


def main() -> None:
    checks: list[dict[str, object]] = []

    def check(identifier: str, condition: bool, detail: str) -> None:
        checks.append({"id": identifier, "ok": bool(condition), "detail": detail})

    required = [
        MANUSCRIPT,
        MANUSCRIPT_BACKUP,
        COVER,
        MODELS / "table3_main_models.csv",
        MODELS / "effect_modification_panel.csv",
        MODELS / "effect_modification_sex_stratified.csv",
        MODELS / "effect_modification_d1_mitml_validation.csv",
        TABLES / "table2_bivariate_observed_applicable.csv",
        SUBMISSION / "S2_Table.xlsx",
        SUBMISSION / "S3_Table.xlsx",
        SUBMISSION / "S7_Table.xlsx",
        REVIEWERS / "REFERENCE_AUDIT_2026-07-23.csv",
        ZENODO_ARCHIVE,
        ZENODO_CHECKSUM,
    ]
    missing = [str(path.relative_to(ROOT)) for path in required if not path.exists()]
    check("inputs", not missing, "Todos los artefactos requeridos existen." if not missing else str(missing))
    if missing:
        raise SystemExit("Faltan entradas: " + ", ".join(missing))

    check(
        "source_backup",
        sha256(MANUSCRIPT_BACKUP) == EXPECTED_SOURCE_SHA,
        f"SHA-256 copia fuente: {sha256(MANUSCRIPT_BACKUP)}",
    )

    main_models = pd.read_csv(MODELS / "table3_main_models.csv")
    main_row = main_models.loc[
        (main_models["model"] == "model_2") & (main_models["term"] == "PHQ9_TOTAL")
    ]
    check("main_model_unique", len(main_row) == 1, f"Filas encontradas: {len(main_row)}")
    main_item = main_row.iloc[0]
    pr_column = "pr" if "pr" in main_row.columns else "exp_estimate"
    lower_column = "ci_low" if "pr" in main_row.columns else "exp_ci_low"
    upper_column = "ci_high" if "pr" in main_row.columns else "exp_ci_high"
    main_pr = float(main_item[pr_column])
    check(
        "main_model_frozen",
        abs(main_pr - 0.9949479) <= 1e-6
        and int(round(float(main_item["mean_n_obs"]))) == 191757,
        (
            f"PR={main_pr:.7f}; IC95%={float(main_item[lower_column]):.7f}–"
            f"{float(main_item[upper_column]):.7f}; p={float(main_item['p_value']):.7f}; "
            f"n={float(main_item['mean_n_obs']):.0f}"
        ),
    )

    panel = pd.read_csv(MODELS / "effect_modification_panel.csv").set_index("modificador")
    expected_modifiers = {"sex", "year", "area", "wealth", "dxhta", "altitude"}
    check("panel_complete", set(panel.index) == expected_modifiers, f"Modificadores: {list(panel.index)}")
    check(
        "modifier_roles",
        panel.loc[["sex", "year"], "tipo"].eq("informado_por_teoria").all()
        and panel.loc[["area", "wealth", "dxhta", "altitude"], "tipo"].eq("exploratorio").all(),
        "Sexo/año informados por teoría; cuatro modificadores exploratorios.",
    )
    check(
        "d1_method",
        panel["pool_method"].eq("rubin_d1_f").all()
        and panel["df_den"].map(math.isfinite).all(),
        "Todos los tests usan rubin_d1_f y df denominador finito.",
    )
    dx = panel.loc["dxhta"]
    check(
        "dx_binary",
        int(dx["df_num"]) == 1
        and int(round(float(dx["mean_n_obs"]))) == 191565
        and "nan" not in str(dx["terms"]).lower(),
        (
            f"df1={int(dx['df_num'])}; n={float(dx['mean_n_obs']):.0f}; "
            f"terms={dx['terms']}"
        ),
    )
    exploratory = panel.loc[panel["tipo"] == "exploratorio"].copy()
    expected_holm = holm(exploratory["p_value"])
    observed_holm = exploratory["p_holm_exploratorios"].astype(float)
    max_holm_error = float((expected_holm - observed_holm).abs().max())
    check("holm", max_holm_error < 1e-12, f"Error absoluto máximo: {max_holm_error:.3g}")

    validation = pd.read_csv(MODELS / "effect_modification_d1_mitml_validation.csv")
    matches = validation["match_1e_10"].map(lambda value: str(value).strip().lower() == "true")
    check(
        "d1_mitml",
        len(validation) == 6 and matches.all(),
        f"{int(matches.sum())}/{len(validation)} tests reproducidos con mitml.",
    )
    slopes = pd.read_csv(MODELS / "effect_modification_sex_stratified.csv")
    sexes = set(slopes["sex"].astype(str).str.lower())
    check(
        "sex_slopes",
        sexes == {"men", "women"}
        and slopes[["pr", "ci_low", "ci_high", "p_value"]].notna().all().all(),
        f"Estratos disponibles: {sorted(sexes)}",
    )

    document = Document(MANUSCRIPT)
    with zipfile.ZipFile(MANUSCRIPT) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml")
    check(
        "no_comments_or_tracked_changes",
        "word/comments.xml" not in names
        and re.search(rb"<w:ins(?:\s|>)", document_xml) is None
        and re.search(rb"<w:del(?:\s|>)", document_xml) is None,
        "Sin comentarios ni cambios controlados residuales.",
    )
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    body = "\n".join(paragraphs)
    table_body = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    lower_body = (body + "\n" + table_body).lower()
    check("document_tables", len(document.tables) == 5, f"Tablas incrustadas: {len(document.tables)}")
    line_numbering = all(
        section._sectPr.find(qn("w:lnNumType")) is not None
        for section in document.sections
    )
    page_fields = all(
        any("PAGE" in paragraph._p.xml for paragraph in section.footer.paragraphs)
        for section in document.sections
    )
    check(
        "line_page_numbering",
        line_numbering and page_fields,
        f"Numeración de líneas={line_numbering}; campos de página={page_fields}.",
    )
    check("abstract_heading", paragraphs.count("Abstract") == 1, f"Encabezados Abstract: {paragraphs.count('Abstract')}")
    abstract_index = paragraphs.index("Abstract")
    next_heading = next(
        (
            index
            for index in range(abstract_index + 1, len(paragraphs))
            if paragraphs[index].lower() == "keywords"
            or paragraphs[index].lower().startswith("keywords:")
        ),
        None,
    )
    abstract_text = " ".join(paragraphs[abstract_index + 1 : next_heading])
    check("abstract_length", word_count(abstract_text) <= 300, f"Resumen: {word_count(abstract_text)} palabras.")
    keyword_lines = [line for line in paragraphs if "keyword" in line.lower()]
    check("keywords_once", len(keyword_lines) == 1, f"Listas/líneas de palabras clave: {len(keyword_lines)}")

    forbidden = [
        "10 mar covariables",
        "rubin_d1_chisq_approx",
        "median pooled non-linearity p",
        "aggregated data",
        "representative cohort",
        "baseline characteristics of the cohort",
        "estimates the total effect",
        "adjusted direct effects",
        "with mediators",
        "protective effect",
        "care paradox was",
        "without 2020 cohort",
        "geographic confounder",
        "10.1097/hjh.0000000000004029",
        "10.1016/j.jad.2021.12.039",
        "hyppertensionaha.110.156307",
    ]
    found = [term for term in forbidden if term in lower_body]
    check("obsolete_text", not found, "Sin términos obsoletos." if not found else str(found))
    required_phrases = [
        "pooled repeated cross-sectional",
        "nine targets were declared",
        "education was not imputed",
        "parametric chained-equations imputation inspired by mice",
        "intimate-partner-violence variable was nominal",
        "mild underdispersion",
        "li–rubin d1 f procedure",
        "the design does not permit causal interpretation",
        "not interpreted as a mediation analysis",
        "s1 table",
        "s4 table",
        "s6 table",
    ]
    absent = [phrase for phrase in required_phrases if phrase not in lower_body]
    if re.search(r"combined(?:\s+across\s+imputations)?\s+using\s+d2", lower_body) is None:
        absent.append("combined [across imputations] using D2")
    check("required_text", not absent, "Todas las frases requeridas están presentes." if not absent else str(absent))
    check("data_statement", DATA_STATEMENT in body, "Declaración de datos exacta presente.")

    reference_start = paragraphs.index("References")
    supporting_start = paragraphs.index("Supporting information")
    reference_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[reference_start + 1 : supporting_start]
        if paragraph.text.strip()
    ]
    references = [paragraph.text.strip() for paragraph in reference_paragraphs]
    check("references_count", len(references) == 47, f"Referencias: {len(references)}")
    numbered_references = [
        paragraph
        for paragraph in reference_paragraphs
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    ]
    check(
        "references_numbered",
        len(numbered_references) == 47,
        f"Referencias con numeración Word: {len(numbered_references)}/47",
    )

    table2_source = pd.read_csv(TABLES / "table2_bivariate_observed_applicable.csv").set_index("label")
    table2 = document.tables[1]
    table2_rows = {
        table2.cell(row, 0).text.strip(): row
        for row in range(1, min(12, len(table2.rows)))
    }
    table2_errors: list[str] = []
    for label, item in table2_source.iterrows():
        if label not in table2_rows:
            table2_errors.append(f"fila ausente: {label}")
            continue
        row = table2_rows[label]
        if table2.cell(row, 3).text.strip() != f"{int(item['n_unweighted']):,}":
            table2_errors.append(f"n discrepante: {label}")
    table2_text = "\n".join(
        cell.text for row in table2.rows for cell in row.cells
    ).lower()
    check(
        "table2_traceability",
        not table2_errors
        and "observed/applicable" in table2_text
        and "not-applicable questionnaire-skip" in table2_text,
        "Tabla 2 concuerda con la salida observada/aplicable." if not table2_errors else str(table2_errors),
    )
    table3_text = "\n".join(
        cell.text for row in document.tables[2].rows for cell in row.cells
    ).lower()
    table4_text = "\n".join(
        cell.text for row in document.tables[3].rows for cell in row.cells
    ).lower()
    check(
        "nine_targets_tables",
        "nine covariates" in table3_text
        and "nine covariates" in table4_text
        and "ten covariates" not in table3_text + table4_text,
        "Tablas 3 y 4 declaran nueve objetivos.",
    )

    s2_text = workbook_text(SUBMISSION / "S2_Table.xlsx").lower()
    s3_text = workbook_text(SUBMISSION / "S3_Table.xlsx").lower()
    s7_text = workbook_text(SUBMISSION / "S7_Table.xlsx").lower()
    check(
        "s2_updated",
        "mild underdispersion" in s2_text
        and "observed" in s2_text
        and "imputed-only pooled" in s2_text
        and "m = 20" in s2_text,
        "S2 contiene φ, diagnóstico observado–imputado y justificación de m.",
    )
    check(
        "s3_updated",
        "d1 li–rubin f" in s3_text
        and "mitml 0.4-5" in s3_text
        and "combined using d2" in s3_text
        and "sex-specific phq-9 slopes" in s3_text
        and "median pooled" not in s3_text,
        "S3 contiene D1 validado, pendientes por sexo y nota D2.",
    )
    check(
        "s7_updated",
        "pooled d1 test" in s7_text
        and "does not by itself prove causal confounding" in s7_text
        and "geographic confounder" not in s7_text,
        "S7 contiene el D1 corregido y lenguaje no causal.",
    )
    all_supplement_text = "\n".join(
        workbook_text(SUBMISSION / f"S{number}_Table.xlsx")
        for number in range(1, 12)
    ).lower()
    supplement_forbidden = [
        "20 mice",
        "with mediators",
        "potential mediators",
        "geographic confounder",
        "median pooled",
        "final analytical cohort",
        "sensitivity cohort without 2020",
    ]
    supplement_found = [
        term for term in supplement_forbidden if term in all_supplement_text
    ]
    check(
        "all_supplements_terminology",
        not supplement_found,
        "Terminología suplementaria sincronizada." if not supplement_found else str(supplement_found),
    )
    figure_details = []
    figure_ok = True
    for filename in ["Fig1.tif", "Fig2.tif", "S1_Fig.tif", "S2_Fig.tif"]:
        with Image.open(SUBMISSION / filename) as image:
            dpi = image.info.get("dpi", (0, 0))
            dpi_values = tuple(float(value) for value in dpi)
            ok = image.mode == "RGB" and min(dpi_values) >= 299
            figure_ok = figure_ok and ok
            figure_details.append(
                f"{filename}:{image.mode}/{dpi_values[0]:.0f}dpi"
            )
    check("figures", figure_ok, "; ".join(figure_details))

    reference_audit = pd.read_csv(REVIEWERS / "REFERENCE_AUDIT_2026-07-23.csv")
    allowed_status = {"verified", "verified_non_doi"}
    bad_references = reference_audit.loc[~reference_audit["status"].isin(allowed_status)]
    check(
        "reference_audit",
        len(reference_audit) == 47 and bad_references.empty,
        (
            "47/47 referencias verificadas."
            if len(reference_audit) == 47 and bad_references.empty
            else f"Referencias={len(reference_audit)}; pendientes={bad_references['reference_number'].tolist()}"
        ),
    )

    zenodo_digest = sha256(ZENODO_ARCHIVE)
    checksum_parts = ZENODO_CHECKSUM.read_text(encoding="ascii").strip().split()
    checksum_ok = (
        len(checksum_parts) == 2
        and checksum_parts[0].upper() == zenodo_digest
        and checksum_parts[1] == ZENODO_ARCHIVE.name
    )
    with zipfile.ZipFile(ZENODO_ARCHIVE) as archive:
        zenodo_names = {
            item.filename
            for item in archive.infolist()
            if not item.is_dir()
        }
        zenodo_bad_suffixes = {
            name
            for name in zenodo_names
            if Path(name).suffix.lower()
            in {
                ".docx",
                ".xlsx",
                ".pdf",
                ".parquet",
                ".dta",
                ".sav",
                ".dbf",
                ".tif",
                ".tiff",
                ".png",
                ".zip",
            }
        }
        required_zenodo_suffixes = {
            "README.md",
            "CITATION.cff",
            "SHA256SUMS.txt",
            "frozen_outputs/analysis/models/effect_modification_panel.csv",
            "frozen_outputs/analysis/models/effect_modification_d1_inputs.json",
            "frozen_outputs/analysis/models/effect_modification_d1_mitml_validation.csv",
            "scripts/audit/validate_effect_modification_d1.R",
        }
        zenodo_paths_ok = all(
            any(name.endswith(suffix) for name in zenodo_names)
            for suffix in required_zenodo_suffixes
        )
        metadata_name = next(
            (
                name
                for name in zenodo_names
                if name.endswith("ZENODO_METADATA_DRAFT_v1.2.0.json")
            ),
            None,
        )
        metadata = (
            json.loads(archive.read(metadata_name).decode("utf-8"))
            if metadata_name
            else {}
        )
    check(
        "zenodo_candidate",
        checksum_ok
        and zenodo_paths_ok
        and not zenodo_bad_suffixes
        and metadata.get("version") == "v1.2.0"
        and metadata.get("access") == "Public",
        (
            f"SHA-256={zenodo_digest}; archivos={len(zenodo_names)}; "
            f"material prohibido={sorted(zenodo_bad_suffixes)}."
        ),
    )

    editorial_text = (SUBMISSION / "PARA_EDITORIAL_MANAGER.md").read_text(encoding="utf-8")
    editorial_normalized = re.sub(r"\s+", " ", editorial_text)
    check(
        "editorial_manager",
        DATA_STATEMENT in editorial_normalized
        and "S3_Fig.tif" not in editorial_text
        and "editorial hold" in lower_body,
        "Formulario sincronizado; marca ética visible.",
    )

    ethics_hold = "editorial hold" in lower_body
    zenodo_release_hold = bool(
        re.search(
            r"La versión pública `v1\.1\.0`\s+corresponde al código anterior a esta ronda",
            editorial_text,
        )
    )
    technical_failures = [item for item in checks if not item["ok"]]
    status = (
        "technical_failure"
        if technical_failures
        else ("external_holds" if ethics_hold or zenodo_release_hold else "complete")
    )
    report = {
        "status": status,
        "manuscript": str(MANUSCRIPT.relative_to(ROOT)).replace("\\", "/"),
        "manuscript_size": MANUSCRIPT.stat().st_size,
        "manuscript_sha256": sha256(MANUSCRIPT),
        "source_backup_sha256": sha256(MANUSCRIPT_BACKUP),
        "ethics_hold": ethics_hold,
        "zenodo_release_hold": zenodo_release_hold,
        "zenodo_candidate_sha256": zenodo_digest,
        "checks": checks,
    }
    REPORT_JSON.write_text(json.dumps(report, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    lines = [
        "# Validación de la guía — 2026-07-23",
        "",
        f"- Estado: **{status}**",
        f"- Manuscrito: `{report['manuscript']}`",
        f"- Tamaño: {report['manuscript_size']:,} bytes",
        f"- SHA-256: `{report['manuscript_sha256']}`",
        "",
        "## Comprobaciones",
        "",
    ]
    for item in checks:
        marker = "x" if item["ok"] else " "
        lines.append(f"- [{marker}] `{item['id']}` — {item['detail']}")
    if ethics_hold:
        lines.extend(
            [
                "",
                "## Bloqueo ético",
                "",
                "El manuscrito conserva `EDITORIAL HOLD`. Falta una determinación institucional "
                "verificable con comité, tipo de decisión, número/fecha y tratamiento del "
                "consentimiento secundario. Este punto no se infiere del acceso público a ENDES.",
            ]
        )
    if zenodo_release_hold:
        lines.extend(
            [
                "",
                "## Publicación del código",
                "",
                "La versión pública `v1.1.0` antecede P0-01/P0-02. Antes del envío debe publicarse "
                "una nueva versión pública bajo el DOI de concepto con el código, el validador D1 y "
                "las salidas congeladas de esta revisión. La API oficial reporta `v1.1.0` con acceso "
                "restringido. El candidato local `v1.2.0` está construido y verificado; aún debe "
                "publicarse y comprobarse mediante descarga anónima.",
            ]
        )
    REPORT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")

    # Mantiene la salida legible también en consolas Windows con cp1252.
    print(json.dumps(report, indent=2, ensure_ascii=True))
    if technical_failures:
        raise SystemExit(1)
    if ethics_hold or zenodo_release_hold:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
