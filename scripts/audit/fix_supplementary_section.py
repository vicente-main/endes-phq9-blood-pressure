# -*- coding: utf-8 -*-
"""Resuelve la inconsistencia S8/S9 en el manuscrito:
- Quita la tabla de operacionalización embebida (su contenido vive en S8_Table.xlsx)
  y su nota de abreviaturas y los captions sueltos S8/S9 del final.
- Crea una sección 'Información complementaria (material suplementario)' con las
  leyendas de TODOS los suplementarios (Tabla S1-S9, Figura S1-S3), solo leyendas,
  porque en PLOS ONE las tablas/figuras suplementarias son archivos separados.

Mantiene español (la traducción y el relabel 'S# Table' son fase posterior).
"""
from __future__ import annotations
from pathlib import Path
import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
DOC = ROOT / "SEVERIDAD SINTOMAS DEPRESIVOS VS PRESION ARTEIAL ELEVADA (1).docx"

TABLAS = [
    ("Tabla S1.", "Flujo de selección muestral STROBE detallado por paso (ENDES 2019-2024)."),
    ("Tabla S2.", "Efectos de diseño (DEFF), diagnósticos de los modelos y detalle de los análisis bivariados por imputación."),
    ("Tabla S3.", "Modelos de sensibilidad e interacción (razones de prevalencia) y modelo logístico de sensibilidad."),
    ("Tabla S4.", "Datos fuente exactos de la Figura 1 (flujo STROBE) y de la Figura S1 (curva spline)."),
    ("Tabla S5.", "Justificación bibliográfica de cada arista del grafo acíclico dirigido (DAG)."),
    ("Tabla S6.", "Comparación de la cohorte incluida frente a los excluidos del análisis (STROBE)."),
    ("Tabla S7.", "Altitud: aporte al estimado (descomposición jerárquica), estratificación y panel de modificación de efecto."),
    ("Tabla S8.", "Operacionalización de las variables del estudio, con su código en los microdatos de la ENDES y su módulo de origen."),
    ("Tabla S9.", "E-values de los estimadores principales (VanderWeele y Ding, 2017): fuerza mínima de la confusión no medida que explicaría la asociación observada."),
]
FIGURAS = [
    ("Figura S1.", "Curva spline cúbica restringida de la prevalencia ajustada de presión arterial elevada según el puntaje PHQ-9, pooled sobre 20 imputaciones MICE (bandas: IC 95 % de Rubin; nodos en 0, 4, 9, 14). La p mediana de no-linealidad fue 0,251 (no se rechaza la linealidad)."),
    ("Figura S2.", "Grafo acíclico dirigido (DAG) de la estructura causal del estudio (exposición PHQ-9 → mediadores → presión arterial elevada), con los confusores del Modelo 2 (incluida la altitud), los mediadores del Modelo 3 y el nodo U de factores no medidos. Verificado acíclico en dagitty."),
    ("Figura S3.", "Forest plot de la razón de prevalencia del PHQ-9 por estrato de altitud (global y categorías < 1 500 / 1 500-2 499 / ≥ 2 500 m s. n. m.), que ilustra la modificación de efecto por altitud (interacción p = 0,017). PR: razón de prevalencia; IC 95 %: intervalo de confianza al 95 %."),
]


def main():
    d = docx.Document(str(DOC))
    body = d.element.body

    # 1) Eliminar tabla de operacionalización + caption S8 + nota abreviaturas + caption S9.
    to_remove = []
    for ch in list(body.iterchildren()):
        if ch.tag == qn("w:tbl"):
            tb = Table(ch, d)
            if len(tb.columns) == 5 and tb.cell(0, 0).text.strip() == "Variable" and len(tb.rows) >= 25:
                to_remove.append(ch)
        elif ch.tag == qn("w:p"):
            t = Paragraph(ch, d).text.strip()
            if (t.startswith("Tabla S8.") or t.startswith("Tabla S9.")
                    or t.startswith("PAS: presión arterial sistólica")):
                to_remove.append(ch)
    for ch in to_remove:
        body.remove(ch)
    print("eliminados del cuerpo:", len(to_remove), "(tabla operacionalización + captions S8/S9 + nota)")

    # 2) Añadir sección de Información complementaria al final.
    h = d.add_paragraph()
    hr = h.add_run("INFORMACIÓN COMPLEMENTARIA (material suplementario)")
    hr.bold = True
    hr.font.size = Pt(12)

    intro = d.add_paragraph()
    ir = intro.add_run("Los siguientes archivos se presentan como información complementaria "
                       "(archivos independientes, no incluidos en el cuerpo del artículo).")
    ir.italic = True
    ir.font.size = Pt(9)

    for label, leg in TABLAS + FIGURAS:
        p = d.add_paragraph()
        rb = p.add_run(label + " ")
        rb.bold = True
        rb.font.size = Pt(10)
        rl = p.add_run(leg)
        rl.font.size = Pt(10)

    d.save(str(DOC))
    # verificación
    d2 = docx.Document(str(DOC))
    print("tablas en el cuerpo ahora:", len(d2.tables), "(esperado 5: Tablas 1-5)")
    full = "\n".join(p.text for p in d2.paragraphs)
    print("¿sección complementaria?:", "INFORMACIÓN COMPLEMENTARIA" in full)
    print("captions Tabla S#:", sum(full.count(l) for l, _ in TABLAS),
          "| Figura S#:", sum(full.count(l) for l, _ in FIGURAS))


if __name__ == "__main__":
    main()
