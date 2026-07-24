"""Reduce y actualiza la cover letter de PLOS ONE a una página real."""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
PATH = ROOT / "ENVIO_PLOS_ONE_2019-2025" / "COVER_LETTER.docx"
BACKUP = PATH.with_name("COVER_LETTER.pre_guia_2026-07-23.docx")
EXPECTED_SHA256 = "4F1D1838FB8ACB4FBB51BA713E025930B1FA6693EE42BD1F5037214C80AE55E8"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    if not BACKUP.exists():
        if sha256(PATH) != EXPECTED_SHA256:
            raise RuntimeError("La cover letter no coincide con la versión auditada.")
        shutil.copy2(PATH, BACKUP)
    if sha256(BACKUP) != EXPECTED_SHA256:
        raise RuntimeError("La copia previa de la cover letter no coincide con la versión auditada.")

    document = Document(BACKUP)
    for paragraph in list(document.paragraphs):
        delete_paragraph(paragraph)

    content = [
        "23 July 2026",
        "To the Editors\nPLOS ONE",
        "Dear Editors,",
        (
            "We submit the Research Article “Severity of depressive symptoms and elevated blood "
            "pressure in Peruvian adults: a cross-sectional analysis of the ENDES 2019–2025” for "
            "consideration in PLOS ONE."
        ),
        (
            "We analyzed de-identified individual-level data from seven annual, nationally "
            "representative Peruvian health surveys (n = 191,757 adults). Survey-weighted "
            "quasi-Poisson models incorporated the complex design, a DAG-informed adjustment set, "
            "20 parametric chained-equations imputations, restricted cubic splines, sensitivity "
            "analyses and validated D1 interaction tests. The main estimate was a very small inverse "
            "cross-sectional association between PHQ-9 score and elevated blood pressure "
            "(PR = 0.995 per point; 95% CI 0.991–0.999). Its clinical importance is uncertain, and "
            "the manuscript avoids causal interpretation. Adding altitude moved the estimate toward "
            "the null, illustrating the methodological importance of geographic context in Andean "
            "population studies."
        ),
        (
            "The work contributes nationally representative evidence from an under-represented Latin "
            "American population and a fully reproducible analysis. Code is archived on Zenodo "
            "(https://doi.org/10.5281/zenodo.21328300), and the ENDES microdata are "
            "publicly available from Peru’s INEI."
        ),
        (
            "This manuscript is original, has not been published, and is not under consideration "
            "elsewhere. All authors approved the submission and declare no competing interests. We "
            "have had no prior interaction with PLOS about this work and do not request exclusion of "
            "any reviewer."
        ),
        (
            "We suggest Dr. Esteban Ortiz-Prado (Universidad de Las Américas, Quito, Ecuador; ORCID "
            "0000-0002-1895-7498) as Academic Editor because of his expertise in cardiovascular and "
            "high-altitude epidemiology. His PLOS ONE editorship was verified in recent publications; "
            "to the best of our knowledge, no author has a conflict of interest with him."
        ),
        "Thank you for your consideration.",
        "Sincerely,",
        (
            "Rodrigo Javier Cardenas Golac, on behalf of all authors\n"
            "Universidad Nacional de la Amazonía Peruana (UNAP), Iquitos, Peru\n"
            "rodrigo.cardenas@unapiquitos.edu.pe"
        ),
    ]

    for index, text in enumerate(content):
        paragraph = document.add_paragraph(text)
        fmt = paragraph.paragraph_format
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.space_after = Pt(4)
        fmt.space_before = Pt(0)
        if index in {2, 9, 10}:
            fmt.space_before = Pt(4)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10.5)

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    document.save(PATH)
    word_count = sum(len(text.replace("\n", " ").split()) for text in content)
    print(f"Cover letter actualizada: {PATH.relative_to(ROOT)}")
    print(f"Palabras: {word_count}")
    print(f"SHA-256 provisional: {sha256(PATH)}")


if __name__ == "__main__":
    main()
