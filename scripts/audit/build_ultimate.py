# -*- coding: utf-8 -*-
"""Crea la versión ULTIMATE = v2 + revisiones de coautores que aún aplican.
Comentarios incorporados: #1, #4, #5, #6, #7-9, #10-11, #15 (ver informe).
Mapa de referencias tras borrar Freedman (#23): 24->23 ... 30->29.
"""
from __future__ import annotations
import copy, shutil
from pathlib import Path
import docx
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
V2 = ROOT / "SEVERIDAD SINTOMAS DEPRESIVOS VS PRESION ARTEIAL ELEVADA v2.docx"
ULT = ROOT / "SEVERIDAD SINTOMAS DEPRESIVOS VS PRESION ARTEIAL ELEVADA ULTIMATE.docx"


def rebuild(p, new_text, lead_bold=None):
    """Reemplaza el contenido del párrafo por new_text en un solo run, copiando el
    formato del primer run. Si lead_bold se da, antepone ese texto en negrita."""
    src = p.runs[0] if p.runs else None
    rpr = copy.deepcopy(src._element.rPr) if (src is not None and src._element.rPr is not None) else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if lead_bold:
        rb = p.add_run(lead_bold)
        rb.bold = True
        if rpr is not None:
            # copiar fuente pero forzar negrita
            rb._element.insert(0, copy.deepcopy(rpr))
            rb.bold = True
    run = p.add_run(new_text)
    if rpr is not None:
        run._element.insert(0, copy.deepcopy(rpr))
        run.bold = False
    return p


def run_replace(p, old, new):
    hit = False
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            hit = True
    return hit


def main():
    shutil.copyfile(V2, ULT)
    d = docx.Document(str(ULT))
    P = d.paragraphs

    # ---- #1: muestreo "de tipo equilibrado" (idx 47) ----
    assert "muestreo probabilístico, bietápico" in P[47].text, P[47].text
    run_replace(P[47], "muestreo probabilístico, bietápico",
                "muestreo probabilístico de tipo equilibrado, bietápico") or rebuild(
        P[47], P[47].text.replace("muestreo probabilístico, bietápico",
                                  "muestreo probabilístico de tipo equilibrado, bietápico"))

    # ---- #4 + #5 + #15(shift diet): reconstruir Covariables (idx 54) ----
    t54 = P[54].text
    assert t54.startswith("Covariables. ")
    body54 = t54[len("Covariables. "):]
    body54 = body54.replace(
        "(HV040, en metros sobre el nivel del mar) La altitud integra",
        "(HV040, en metros sobre el nivel del mar), y se modeló de forma categórica "
        "(< 1 500, 1 500–2 499 y ≥ 2 500 m s. n. m.) (26). La altitud integra")
    body54 = body54.replace("confusor geográfico de la asociación (26, 27) y, si bien",
                            "confusor geográfico de la asociación y, si bien")          # #5
    body54 = body54.replace("Organización Mundial de la Salud (24, 25)—",
                            "Organización Mundial de la Salud (23, 24)—")                # #15 shift
    assert "(26, 27)" not in body54 and "(26)" in body54 and "(23, 24)" in body54
    rebuild(P[54], body54, lead_bold="Covariables. ")

    # ---- #6: quitar frase redundante de Rubin en Datos faltantes (idx 57) ----
    assert P[57].text.rstrip().endswith("se combinaron mediante las reglas de Rubin.")
    t57 = P[57].text.replace(
        " Las estimaciones de los conjuntos imputados se combinaron mediante las reglas de Rubin.", "")
    rebuild(P[57], t57.replace("Covariables. ", "", 0))  # no lead bold aquí
    # (P[57] no empieza con 'Covariables.'; rebuild sin lead)

    # ---- #7-9: reescritura de Análisis estadístico (idx 59) ----
    stats = (
        "Los análisis incorporaron el diseño muestral complejo de la ENDES mediante el "
        "paquete survey (27) de R (versión 4.5.3), declarando el conglomerado primario "
        "(HV001), el estrato (HV022) y el factor de ponderación analítico multianual con la "
        "función svydesign (con la opción nest = TRUE). El peso final se obtuvo a partir del "
        "peso de salud original (PESO15_AMAS) dividido entre un millón y de-normalizado por el "
        "número de años combinados (k = 6). Las varianzas se estimaron por linealización de "
        "Taylor; para los estratos con un único conglomerado activo (singleton strata) se "
        "aplicó el ajuste centrado en la gran media (opción survey.lonely.psu = «adjust»). En "
        "cumplimiento de las directrices STROBE para encuestas poblacionales, se calculó el "
        "efecto de diseño (DEFF) de las proporciones principales. El análisis inferencial se "
        "ejecutó de forma independiente en cada uno de los 20 conjuntos imputados, incorporando "
        "el diseño complejo, y los estimadores se consolidaron mediante las reglas de Rubin (28)."
    )
    assert "inflación artificial" in P[59].text and "survey (28)" in P[59].text
    rebuild(P[59], stats)

    # ---- #10-11: acortar bivariado (idx 60) ----
    assert "su extensión apropiada para muestras de diseño complejo" in P[60].text
    t60 = P[60].text.replace(
        "con corrección de Rao-Scott, su extensión apropiada para muestras de diseño complejo.",
        "con corrección de Rao-Scott.")
    rebuild(P[60], t60)

    # ---- #15 shift: VanderWeele (30)->(29) en métodos E-values (idx 63) ----
    assert "(30)" in P[63].text
    run_replace(P[63], "(30)", "(29)") or rebuild(P[63], P[63].text.replace("(30)", "(29)"))

    # ---- #15: borrar la referencia Freedman de la lista ----
    removed = False
    for p in d.paragraphs:
        if p.text.strip().startswith("Freedman DS, Lawman HG"):
            p._element.getparent().remove(p._element)
            removed = True
            break
    assert removed, "No se encontró la referencia Freedman"

    d.save(str(ULT))

    # ---- verificación ----
    d2 = docx.Document(str(ULT))
    body = []
    refs = False
    import re
    for p in d2.paragraphs:
        if p.text.strip().startswith("10. REFERENCIAS"):
            refs = True
        if not refs:
            body.append(p.text)
    full = "\n".join(body)
    cites = sorted(set(re.findall(r"\([0-9]{1,2}(?:\s*,\s*[0-9]{1,2})*\)", full)))
    print("OK ->", ULT.name)
    print("Citas in-text (cuerpo):", cites)
    print("Freedman en lista:", any(p.text.strip().startswith("Freedman DS") for p in d2.paragraphs))
    print("'equilibrado':", "de tipo equilibrado" in full)
    print("'inflación artificial' (debe ser False):", "inflación artificial" in full)
    print("'su extensión apropiada' (debe ser False):", "su extensión apropiada" in full)
    print("categórica altitud:", "se modeló de forma categórica" in full)
    print("tablas en cuerpo:", len(d2.tables))


if __name__ == "__main__":
    main()
