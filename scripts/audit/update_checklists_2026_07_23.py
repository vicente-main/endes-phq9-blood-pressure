"""Sincroniza las referencias de ubicación en los checklists STROBE/RECORD."""

from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SUBMISSION = ROOT / "ENVIO_PLOS_ONE_2019-2025"
FILES = {
    SUBMISSION / "S1_Checklist.docx": (
        "9D32B45908D2539480E9AD3C8189C36216CBF092BEC0B952195A044BE096F406",
        {
            "Missing-data handling (MICE, m=20); survey design (svydesign); "
            "effect-modification panel; S2/S3/S7 Tables": (
                "Missing-data handling (20 parametric chained-equations imputations); "
                "survey design (svydesign); D1-validated effect-modification panel and "
                "sex-specific estimates; S2/S3/S7 Tables"
            ),
            "Variables; S8 Table; DAG (S2 Fig, S5 Table)": (
                "Variables; S8 Table; DAG (S1 Fig, S5 Table)"
            ),
            "Results; Table 5; S3 Table (sensitivity, effect modification); "
            "S7 Table (altitude); S1 Fig (spline); S3 Fig (altitude forest)": (
                "Results; Table 5; S3 Table (sensitivity and effect modification); "
                "S7 Table (altitude); Fig 2 (spline); S2 Fig (altitude forest)"
            ),
        },
    ),
    SUBMISSION / "S2_Checklist.docx": (
        "C85646C8AF4A4DDF52A2B3E346DDAD49BF2E1081A09A309EF3355372D4635E21",
        {
            "MICE (m = 20) for MAR covariables.": (
                "20 parametric chained-equations imputations for covariates under MAR, "
                "with observed-versus-imputed diagnostics in S2 Table."
            ),
            "covariate roles are justified edge-by-edge in the DAG (S2 Fig).": (
                "covariate roles are justified edge-by-edge in the DAG (S1 Fig)."
            )
        },
    ),
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def main() -> None:
    for path, (expected_hash, replacements) in FILES.items():
        backup = path.with_name(f"{path.stem}.pre_guia_2026-07-23{path.suffix}")
        if not backup.exists():
            if sha256(path) != expected_hash:
                raise RuntimeError(f"{path.name} no coincide con la versión auditada.")
            shutil.copy2(path, backup)
        if sha256(backup) != expected_hash:
            raise RuntimeError(f"La copia fuente de {path.name} no coincide con la versión auditada.")

        document = Document(backup)
        counts = {old: 0 for old in replacements}
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        for old, new in replacements.items():
                            if old in text:
                                paragraph.text = text.replace(old, new)
                                counts[old] += 1
                                text = paragraph.text
        missing = [old for old, count in counts.items() if count == 0]
        if missing:
            raise RuntimeError(f"No se encontraron textos esperados en {path.name}: {missing}")
        document.save(path)
        print(f"Actualizado: {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
