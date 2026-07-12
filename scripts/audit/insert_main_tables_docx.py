# -*- coding: utf-8 -*-
"""Inserta las Tablas 1-5 (principales) como tablas Word EDITABLES dentro del
manuscrito, justo tras su primera cita (requisito PLOS ONE: tablas en el .docx,
basadas en celdas, no como imagen). Mantiene el contenido en español (la
traducción es una fase posterior). Replica combinaciones de celda y sombreados.

Fuente: data/output/Post_Auditoria/Principal/Tabla_{1..5}.xlsx (hoja de datos).
"""
from __future__ import annotations

import copy
from pathlib import Path

import docx
import openpyxl
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
DOC = ROOT / "SEVERIDAD SINTOMAS DEPRESIVOS VS PRESION ARTEIAL ELEVADA (1).docx"
PRIN = ROOT / "data" / "output" / "Post_Auditoria" / "Principal"

HEAD_FILL = "D9E1F2"
SEC_FILL = "F2F2F2"

# Tabla -> (n_filas_encabezado, párrafo-ancla tras el cual insertar)
# Tabla 1 y 2 se citan primero en el mismo párrafo de Resultados (idx 68).
PLAN = [
    ("Tabla_1", 2, 68),
    ("Tabla_2", 1, 68),
    ("Tabla_3", 1, 70),
    ("Tabla_4", 1, 72),
    ("Tabla_5", 1, 74),
]


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tblPr.append(borders)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text):
    """Fija el texto de la celda en UN solo párrafo, sin párrafos vacíos residuales
    (evita las líneas en blanco que deja .merge()). '\\n' -> salto de línea."""
    # vaciar todos los párrafos menos el primero
    for p in list(cell.paragraphs[1:]):
        p._p.getparent().remove(p._p)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._r.getparent().remove(r._r)
    parts = str(text).split("\n")
    run = p0.add_run(parts[0])
    for extra in parts[1:]:
        run.add_break()
        run = p0.add_run(extra)


def style_cell(cell, *, bold=False, size=9, align="left", fill=None, italic=False):
    for p in cell.paragraphs:
        p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
        for run in p.runs:
            run.font.bold = bold
            run.font.italic = italic
            run.font.size = Pt(size)
    if fill:
        set_cell_shading(cell, fill)


def read_sheet(path):
    wb = openpyxl.load_workbook(path)
    name = [s for s in wb.sheetnames if s != "Portada"]
    ws = wb[name[0]] if name else wb.active
    nrow, ncol = ws.max_row, ws.max_column

    def val(r, c):
        v = ws.cell(r, c).value
        return "" if v is None else str(v).strip()

    title = val(1, 1)
    # filas de datos = 2..nrow
    body = [[val(r, c) for c in range(1, ncol + 1)] for r in range(2, nrow + 1)]
    # merges en filas >= 2, trasladados a coords de la tabla Word (0-indexed, sin la fila título)
    merges = []
    for rng in ws.merged_cells.ranges:
        if rng.min_row >= 2:
            merges.append((rng.min_row - 2, rng.min_col - 1, rng.max_row - 2, rng.max_col - 1))
    return title, body, ncol, merges


def is_full_width_row(body_row):
    nonempty = [x for x in body_row if x]
    return len(nonempty) == 1 and body_row[0] != ""


def classify_footer(text):
    t = text.strip()
    low = t.lower()
    if low.startswith(("fuente", "pas:", "pad:", "abreviat", "nota:")) or "fuente:" in low:
        return "footer"
    if low.startswith(("lectura", "interpretaci")):
        return "note"
    return "section"


def make_paragraph_after(anchor_p, text, *, bold=False, size=11):
    new_p = OxmlElement("w:p")
    anchor_p._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, anchor_p._parent)
    run = para.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    return para


def build_and_insert(doc, anchor_p, title, body, ncol, n_header, merges):
    nrow = len(body)
    table = doc.add_table(rows=nrow, cols=ncol)
    set_table_borders(table)
    table.autofit = True

    # Lista de combinaciones a aplicar: las del xlsx + las filas de ancho completo
    # (secciones/notas/pie) que en el origen quedaron como celdas sueltas.
    applied = [(r1, c1, r2, c2) for (r1, c1, r2, c2) in merges
               if 0 <= r1 < nrow and 0 <= r2 < nrow and 0 <= c1 < ncol and 0 <= c2 < ncol]
    for r in range(nrow):
        if r >= n_header and is_full_width_row(body[r]):
            applied.append((r, 0, r, ncol - 1))

    covered = set()
    for (r1, c1, r2, c2) in applied:
        try:
            table.cell(r1, c1).merge(table.cell(r2, c2))
        except Exception:
            pass
        for rr in range(r1, r2 + 1):
            for cc in range(c1, c2 + 1):
                if (rr, cc) != (r1, c1):
                    covered.add((rr, cc))

    # Rellenar texto solo en las celdas maestras (las combinadas heredan del maestro)
    for r in range(nrow):
        for c in range(ncol):
            if (r, c) in covered:
                continue
            set_cell_text(table.cell(r, c), body[r][c])

    # Estilos
    for r in range(nrow):
        row = body[r]
        if r < n_header:
            for c in range(ncol):
                if (r, c) in covered:
                    continue
                style_cell(table.cell(r, c), bold=True, size=9,
                           align="center" if c > 0 else "left", fill=HEAD_FILL)
        elif is_full_width_row(row):
            kind = classify_footer(row[0])
            if kind == "footer":
                style_cell(table.cell(r, 0), bold=False, italic=True, size=8, align="left")
            elif kind == "note":
                style_cell(table.cell(r, 0), bold=False, italic=True, size=9, align="left")
            else:
                style_cell(table.cell(r, 0), bold=True, size=9, align="left", fill=SEC_FILL)
        else:
            for c in range(ncol):
                if (r, c) in covered:
                    continue
                style_cell(table.cell(r, c), bold=False, size=9,
                           align="left" if c == 0 else "center")

    # Caption (leyenda) + tabla tras el ancla: ancla -> caption -> tabla
    tbl_el = table._tbl
    tbl_el.getparent().remove(tbl_el)          # quitar del final del doc
    anchor_p._p.addnext(tbl_el)                # insertar tras el ancla
    make_paragraph_after(anchor_p, title, bold=True, size=10)  # caption antes de la tabla
    return table


def main():
    doc = docx.Document(str(DOC))
    paras = doc.paragraphs
    # Insertar en orden inverso de documento para no invalidar anclas;
    # para el ancla compartida (68) se inserta Tabla_2 antes que Tabla_1 -> queda T1, T2.
    for name, n_header, anchor_idx in reversed(PLAN):
        title, body, ncol, merges = read_sheet(PRIN / f"{name}.xlsx")
        anchor_p = paras[anchor_idx]
        build_and_insert(doc, anchor_p, title, body, ncol, n_header, merges)
        print(f"insertada {name}: {len(body)} filas x {ncol} cols tras idx {anchor_idx}")

    doc.save(str(DOC))
    d2 = docx.Document(str(DOC))
    print("Tablas en el documento ahora:", len(d2.tables))


if __name__ == "__main__":
    main()
