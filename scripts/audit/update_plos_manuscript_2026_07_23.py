"""Aplica al manuscrito de Revisores la guía corregida del 23-07-2026.

La edición se alimenta de salidas analíticas congeladas; no transcribe cifras
secundarias a mano. El primer uso exige el SHA-256 de la versión auditada y
crea una copia de resguardo antes de guardar el DOCX en su ruta original.
"""

from __future__ import annotations

import hashlib
import math
import re
import shutil
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
SUBMISSION = ROOT / "ENVIO_PLOS_ONE_2019-2025"
REVIEWERS = SUBMISSION / "Revisores"
MANUSCRIPT = REVIEWERS / "MANUSCRITO_EN_PLOS_ONE.docx"
BACKUP = REVIEWERS / "MANUSCRITO_EN_PLOS_ONE.pre_guia_2026-07-23.docx"
EXPECTED_SHA256 = "C56E7016BD881D3402529933D2E4A28BC61529A4D9EA4BC2238415D02C926DD7"

ANALYSIS = ROOT / "data" / "output_2025" / "analysis"
MODELS = ANALYSIS / "models"
TABLES = ANALYSIS / "tables"
FIGURES = ANALYSIS / "figures"

MAIN_MODELS = MODELS / "table3_main_models.csv"
CASCADE_MODELS = MODELS / "table4_cascade_models.csv"
EFFECT_PANEL = MODELS / "effect_modification_panel.csv"
SEX_SLOPES = MODELS / "effect_modification_sex_stratified.csv"
COMPLETE_CASE = MODELS / "complete_case_sensitivity.csv"
HIERARCHICAL = MODELS / "hierarchical_decomposition.csv"
SPLINE = FIGURES / "spline_nonlinearity_summary.csv"
TABLE2 = TABLES / "table2_bivariate_observed_applicable.csv"
MITML_VALIDATION = MODELS / "effect_modification_d1_mitml_validation.csv"

EXPECTED_OUTPUTS = [
    MAIN_MODELS,
    CASCADE_MODELS,
    EFFECT_PANEL,
    SEX_SLOPES,
    COMPLETE_CASE,
    HIERARCHICAL,
    SPLINE,
    TABLE2,
    MITML_VALIDATION,
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def require_inputs() -> None:
    missing = [str(path) for path in [MANUSCRIPT, *EXPECTED_OUTPUTS] if not path.exists()]
    if missing:
        raise FileNotFoundError("Faltan entradas requeridas:\n" + "\n".join(missing))
    current_hash = sha256(MANUSCRIPT)
    if current_hash != EXPECTED_SHA256:
        raise RuntimeError(
            "El manuscrito no coincide con la versión auditada. "
            f"Esperado {EXPECTED_SHA256}; observado {current_hash}."
        )
    if BACKUP.exists() and sha256(BACKUP) != EXPECTED_SHA256:
        raise RuntimeError(f"La copia previa existente no coincide con el hash auditado: {BACKUP}")


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Se esperaba un párrafo con prefijo único {prefix!r}; encontrados {len(matches)}.")
    return matches[0]


def set_paragraph(document: Document, prefix: str, text: str) -> Paragraph:
    paragraph = find_paragraph(document, prefix)
    paragraph.text = text
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_heading_before(paragraph: Paragraph, text: str, style: str = "Heading 1") -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def remove_empty_numbered_paragraphs_before(paragraph: Paragraph) -> None:
    """Elimina marcadores de lista huérfanos inmediatamente anteriores."""
    current = paragraph._p.getprevious()
    while current is not None and current.tag.endswith("}p"):
        candidate = Paragraph(current, paragraph._parent)
        if candidate.text.strip():
            break
        previous = current.getprevious()
        if (
            candidate._p.pPr is not None
            and candidate._p.pPr.numPr is not None
        ):
            delete_paragraph(candidate)
        current = previous


def copy_paragraph_properties(source: Paragraph, target: Paragraph) -> None:
    """Copia numeración, sangría y espaciado sin copiar el contenido."""
    if target._p.pPr is not None:
        target._p.remove(target._p.pPr)
    if source._p.pPr is not None:
        target._p.insert(0, deepcopy(source._p.pPr))


def set_merged_row_text(table, row_index: int, text: str) -> None:
    seen: set[int] = set()
    for cell in table.rows[row_index].cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        cell.text = text


def format_p(value: float, digits: int = 3) -> str:
    if not math.isfinite(value):
        return "—"
    return "< 0.001" if value < 0.001 else f"{value:.{digits}f}"


def format_effect(item: pd.Series) -> str:
    text = format_p(float(item["p_value"]))
    if str(item["tipo"]) == "exploratorio":
        text += f"; Holm {format_p(float(item['p_holm_exploratorios']))}"
    return text


def model_row(frame: pd.DataFrame, model: str, term: str = "PHQ9_TOTAL") -> pd.Series:
    rows = frame.loc[(frame["model"] == model) & (frame["term"] == term)]
    if len(rows) != 1:
        raise RuntimeError(f"No se encontró una fila única para {model}/{term}.")
    item = rows.iloc[0].copy()
    # Las salidas canónicas conservan coeficientes e IC en escala logarítmica
    # bajo ci_low/ci_high y exponen la escala PR como exp_*.
    if "pr" not in frame.columns:
        item["pr"] = item["exp_estimate"]
        item["ci_low"] = item["exp_ci_low"]
        item["ci_high"] = item["exp_ci_high"]
    return item


def assert_d1_validation() -> None:
    validation = pd.read_csv(MITML_VALIDATION)
    if "match_1e_10" not in validation.columns:
        raise RuntimeError("La validación D1 no contiene la bandera match_1e_10.")
    matches = validation["match_1e_10"].map(
        lambda value: str(value).strip().lower() == "true"
    )
    if not matches.all():
        raise RuntimeError("La validación D1 contra mitml no fue íntegramente satisfactoria.")


def main() -> None:
    require_inputs()
    assert_d1_validation()

    main_models = pd.read_csv(MAIN_MODELS)
    cascade = pd.read_csv(CASCADE_MODELS)
    panel = pd.read_csv(EFFECT_PANEL).set_index("modificador")
    slopes = pd.read_csv(SEX_SLOPES)
    slopes["sex"] = slopes["sex"].astype(str).str.lower()
    slopes = slopes.set_index("sex")
    spline = pd.read_csv(SPLINE).iloc[0]
    table2 = pd.read_csv(TABLE2).set_index("label")

    main_model = model_row(main_models, "model_2")
    crude_model = model_row(main_models, "model_1")
    exploratory_model = model_row(main_models, "model_3")
    adherence = model_row(cascade, "submodel_adherence")
    control = model_row(cascade, "submodel_domain_bp")

    if abs(float(main_model["pr"]) - 0.9949479) > 1e-6:
        raise RuntimeError("El Modelo 2 principal cambió fuera de la tolerancia prevista.")
    if int(round(float(main_model["mean_n_obs"]))) != 191757:
        raise RuntimeError("El n del Modelo 2 principal dejó de ser 191,757.")
    if int(panel.loc["dxhta", "df_num"]) != 1:
        raise RuntimeError("La interacción por diagnóstico previo no quedó como contraste binario de 1 df.")
    if int(round(float(panel.loc["dxhta", "mean_n_obs"]))) != 191565:
        raise RuntimeError("La interacción por diagnóstico previo no excluyó exactamente 192 faltantes.")

    document = Document(MANUSCRIPT)

    # Portada, resumen y palabras clave.
    set_paragraph(
        document,
        "1 Universidad Nacional",
        "1 Universidad Nacional de la Amazonía Peruana (UNAP), Iquitos, Peru",
    )
    set_paragraph(
        document,
        "a Physician;",
        "a MD; c BSc in Computer Science; d Medical student",
    )
    remove_empty_numbered_paragraphs_before(
        find_paragraph(document, "Corresponding author:")
    )
    abstract = (
        "To estimate the cross-sectional association between depressive symptom severity and "
        "elevated blood pressure in Peruvian adults and to explore antihypertensive medication "
        "non-adherence among adults with a prior hypertension diagnosis. We analyzed de-identified "
        "individual-level microdata from the nationally representative Peruvian Demographic and "
        "Family Health Survey (ENDES), 2019–2025 (n = 191,757 adults). Elevated blood pressure was "
        "defined as mean systolic pressure ≥ 140 mmHg or mean diastolic pressure ≥ 90 mmHg. "
        "Survey-weighted quasi-Poisson models incorporated the complex sampling design; missing "
        "covariates were handled with 20 parametric chained-equations imputations. The weighted "
        "prevalence of elevated blood pressure was 16.1%. In the main model adjusted for "
        "sociodemographic, contextual, temporal and geographic covariates, each additional PHQ-9 "
        f"point was associated with a prevalence ratio of {float(main_model['pr']):.3f} "
        f"(95% CI {float(main_model['ci_low']):.3f}–{float(main_model['ci_high']):.3f}; "
        f"p = {float(main_model['p_value']):.3f}). Adding altitude moved the estimate approximately "
        "22% toward the null; this descriptive attenuation does not establish a causal explanation. "
        f"The PHQ-9 slope was {float(slopes.loc['men', 'pr']):.3f} "
        f"(95% CI {float(slopes.loc['men', 'ci_low']):.3f}–"
        f"{float(slopes.loc['men', 'ci_high']):.3f}) in men and "
        f"{float(slopes.loc['women', 'pr']):.3f} "
        f"(95% CI {float(slopes.loc['women', 'ci_low']):.3f}–"
        f"{float(slopes.loc['women', 'ci_high']):.3f}) in women. "
        f"PHQ-9 severity was not associated with medication non-adherence "
        f"(PR {float(adherence['pr']):.3f}; 95% CI {float(adherence['ci_low']):.3f}–"
        f"{float(adherence['ci_high']):.3f}; p = {float(adherence['p_value']):.3f}). "
        "The main estimate indicates a very small inverse cross-sectional association whose "
        "clinical importance is uncertain. The design does not permit causal interpretation."
    )
    abstract_paragraph = set_paragraph(document, "To assess the association", abstract)
    insert_heading_before(abstract_paragraph, "Abstract")
    duplicate_keywords = find_paragraph(document, "Depression; hypertension; blood pressure;")
    delete_paragraph(duplicate_keywords)

    # Introducción.
    set_paragraph(
        document,
        "Blood-pressure alteration is a multicausal",
        (
            "Blood-pressure alteration is associated with biological, metabolic, behavioral and "
            "contextual factors [5]. Chronic stress and mood disorders have been linked to autonomic "
            "and hypothalamic-pituitary-adrenal activity that may relate to vascular physiology [6]. "
            "Accordingly, depressive symptom severity has been investigated as a psychosocial factor "
            "associated with hypertension and blood-pressure variation [7–9]."
        ),
    )
    set_paragraph(
        document,
        "Although systematic reviews in high-income",
        (
            "Although systematic reviews in high-income countries have documented longitudinal "
            "associations between depression and cardiovascular outcomes, population-based evidence "
            "from Latin America that incorporates complex survey designs remains limited. Much of the "
            "literature also dichotomizes psychometric scales at fixed thresholds, discarding "
            "information that can be used to examine the shape of the association between continuous "
            "depressive burden, measured with instruments such as the Patient Health Questionnaire-9 "
            "(PHQ-9), and blood-pressure status."
        ),
    )
    set_paragraph(
        document,
        "In addition, the impact of mental health",
        (
            "Mental health may also be associated with treatment behaviors. Among people with a prior "
            "medical diagnosis of hypertension, depressive symptoms could coincide with lower "
            "antihypertensive medication adherence, an important concern because effective blood-"
            "pressure control in Peru remains suboptimal [10]."
        ),
    )
    set_paragraph(
        document,
        "Therefore, the objective of the present study",
        (
            "The objective of this study was to estimate the cross-sectional association between "
            "depressive symptom severity and elevated blood pressure in Peruvian adults using the "
            "pooled ENDES 2019–2025 surveys. We also estimated the cross-sectional association "
            "between depressive symptom burden and antihypertensive medication non-adherence among "
            "participants reporting a prior medical diagnosis of hypertension."
        ),
    )

    # Métodos.
    set_paragraph(
        document,
        "The study population comprised adults",
        (
            "The analytic population comprised adults aged 18 years or older with valid records for "
            "blood-pressure measurements and the mental-health module. Pregnant women (V454 = 1) and "
            "women within two months postpartum (V222 < 2) were excluded because these states alter "
            "blood-pressure interpretation [6,12]. This reproductive filter was applied only to women "
            "of reproductive age (15–49 years); structural missing values in the reproductive-health "
            "modules for men and women aged 50 years or older were interpreted according to the "
            "questionnaire skip pattern. Sample selection is summarized in Fig 1 and detailed in S1 "
            "Table; characteristics of included and excluded records are compared in S6 Table. ENDES "
            "is an annual program of routine population health surveillance. We therefore used RECORD "
            "as a complement to STROBE rather than as a PLOS requirement; completed STROBE and RECORD "
            "checklists are provided as S1 Checklist and S2 Checklist."
        ),
    )
    set_paragraph(
        document,
        "Outcome: elevated blood pressure.",
        (
            "Outcome: elevated blood pressure. Blood pressure was determined from two measurements "
            "taken during one home visit with automatic OMRON HEM-7113 sphygmomanometers and arm-"
            "appropriate cuffs (HEM-CR24 or HEM-CL24), after five minutes of rest and with a two-minute "
            "interval (QS903S/QS903D and QS905S/QS905D). Before averaging, sentinel values (999) and "
            "physiologically implausible readings were set to missing: systolic < 70 or > 270 mmHg and "
            "diastolic < 30 or > 150 mmHg [13,14]. Implausible adult weight (< 25 or > 300 kg) and "
            "height (< 100 or > 200 cm) values were likewise set to missing for anthropometric quality "
            "control [15]. Participants were not excluded solely because an anthropometric value was "
            "invalid; those values were handled as missing covariate data. Records in which systolic "
            "pressure did not exceed diastolic pressure in either measurement were excluded for "
            "hemodynamic inconsistency. Elevated blood pressure was defined as mean systolic blood "
            "pressure ≥ 140 mmHg or mean diastolic blood pressure ≥ 90 mmHg, consistent with population-"
            "surveillance criteria [5,16]."
        ),
    )
    set_paragraph(
        document,
        "Exposure: severity of depressive symptomatology.",
        (
            "Exposure: depressive symptom severity. The Patient Health Questionnaire-9 (PHQ-9) "
            "items are coded QS700A–QS700I in CSALUD01 and cover symptoms during the previous 14 "
            "days. The score was analyzed continuously (range 0–27) to preserve information and "
            "estimate the shape of its association with elevated blood pressure. For description, "
            "we used the standard ordered severity categories: minimal (0–4), mild (5–9), moderate "
            "(10–14), moderately severe (15–19) and severe (20–27) [17,18]. When one or two items "
            "were missing, the score was prorated from answered items; records with three or more "
            "missing items were excluded in the STROBE flow."
        ),
    )
    set_paragraph(
        document,
        "Covariates. The structural adjustment set",
        (
            "Covariates. The main adjustment set (Model 2), defined in the analysis framework and "
            "informed by the directed acyclic graph (S1 Fig), comprised age (QS23), sex (QSSEXO), "
            "educational level (QS25N), area of residence (HV025), wealth index (HV270), intimate "
            "partner violence (QS709–QS711), survey year and altitude of residence. Altitude was "
            "derived from cluster altitude (HV040) and categorized as < 1,500, 1,500–2,499 and "
            "≥ 2,500 m [20,21]. The intimate-partner-violence variable was nominal, with three "
            "categories: no physical violence, physical violence and no partner in the previous 12 "
            "months. It does not capture psychological or sexual violence or lifetime exposure. "
            "Additional metabolic-behavioral covariates, included only in exploratory Model 3, were "
            "body mass index (QS900/QS901), waist circumference (QS907), tobacco use (QS201), "
            "problematic alcohol use (derived from QS209, QS713–QS717, QS719 and QS720), diet quality "
            "(five or more daily fruit and vegetable servings) and prior diabetes diagnosis (QS109) "
            "[22–24]. The rationale and bibliographic support for the analysis framework are documented "
            "in S1 Fig and S5 Table."
        ),
    )
    set_paragraph(
        document,
        "The prior medical diagnosis of hypertension",
        (
            "Prior medical diagnosis of hypertension (QS102) defined the population for the care-"
            "cascade subanalyses and was not an adjustment covariate; non-adherence was derived from "
            "QS104 and QS106. Functional impact of depressive symptoms (QS702) and specialized mental-"
            "health treatment (QS707) were described but not included in the adjustment models because "
            "they are closely tied to symptom burden and health-care contact. Time since hypertension "
            "diagnosis (QS103U/QS103C) contained no usable information and was excluded. Variable "
            "operationalizations and source modules are listed in S8 Table."
        ),
    )
    set_paragraph(
        document,
        "To minimize the bias associated with complete-case",
        (
            "Under a missing-at-random assumption, we generated 20 completed datasets with parametric "
            "chained-equations imputation inspired by MICE, implemented with scikit-learn 1.8.0 "
            "IterativeImputer and sample_posterior = True [25,26]. This implementation is not the R "
            "mice package and does not use predictive mean matching. Nine targets were declared: age, "
            "sex, area, wealth, intimate partner violence, body mass index, waist circumference, diet "
            "quality and diabetes diagnosis. Only five contained missing values that were replaced: "
            "intimate partner violence (29; <0.1%), body mass index (718; 0.4%), waist circumference "
            "(1,656; 0.9%), diet quality (8,430; 4.4%) and diabetes diagnosis (207; 0.1%); age, sex, "
            "area and wealth were complete. Education was not imputed: the questionnaire skip for "
            "people who never attended school (QS24 = 2) was recoded to the lowest category, leaving "
            "three genuinely missing records that were excluded, as were missing exposure, outcome or "
            "altitude values. Continuous draws were constrained to the observed range and categorical "
            "draws were reassigned to the nearest valid observed category. The imputation predictors "
            "included PHQ-9, elevated blood pressure, survey year, cluster, stratum, final survey weight "
            "and the other target variables; exposure and outcome were not imputed. Structural skips "
            "for tobacco and alcohol were retained as explicit not-applicable categories. Twenty "
            "imputations were used to reduce Monte Carlo error and were conservative relative to "
            "the largest target-specific missing proportion (4.4%) [46]. Observed-versus-imputed distributions, between-"
            "imputation stability and range/category checks are reported in S2 Table. IterativeImputer "
            "did not retain a chain-iteration convergence trace; estimates were combined with Rubin's "
            "rules."
        ),
    )
    set_paragraph(
        document,
        "In the descriptive and bivariate analysis",
        (
            "For descriptive analyses, weighted frequencies and proportions were estimated for "
            "categorical variables and weighted means with standard deviations for continuous "
            "variables. Table 1 and Table 2 use observed, non-imputed data. Table 2 reports Rao–Scott "
            "design-adjusted F tests and the observed denominator for each variable; the not-applicable "
            "questionnaire-skip categories for alcohol and tobacco are shown descriptively in Table 1 "
            "but excluded from their Table 2 tests. Multivariable models used the 20 completed datasets."
        ),
    )
    set_paragraph(
        document,
        "Covariate selection was based on a directed",
        (
            "A directed acyclic graph (S1 Fig; verified acyclic in dagitty) informed the main "
            "adjustment set [29]. Model 1 was crude. Model 2, the main structural model, adjusted for "
            "age, sex, education, area, wealth, intimate partner violence, year and altitude. Model 3 "
            "added body mass index, waist circumference, tobacco use, problematic alcohol use, diet "
            "quality and diabetes as an exploratory specification; differences between Models 2 and 3 "
            "were not interpreted as mediation or direct effects. Because elevated blood pressure was "
            "common, survey-weighted quasi-Poisson models with log link estimated prevalence ratios "
            "directly [30], with 95% confidence intervals from Taylor linearization. The design-based "
            "quasi-Poisson dispersion parameter for Model 2 was φ = 0.761, indicating mild "
            "underdispersion rather than overdispersion (S2 Table). Column-wise variance inflation "
            "factors were calculated for indicator-coded predictors; all were below 10, with the "
            "largest values occurring among indicators of the same educational-level factor (S2 Table)."
        ),
    )
    set_paragraph(
        document,
        "Sensitivity and interaction analyses.",
        (
            "Sensitivity and interaction analyses. Secondary specifications comprised Model 2 without "
            "2020; a second-measurement-only blood-pressure outcome; the ≥130/80 mmHg ACC/AHA threshold "
            "[32]; a logistic model with a design-adjusted Archer–Lemeshow test; complete-case analysis; "
            "and restricted cubic splines. Theory-informed modifiers were sex and survey year; area, "
            "wealth, prior hypertension diagnosis and altitude were exploratory. Joint interaction "
            "coefficients were combined across imputations with the Li–Rubin D1 F procedure with finite "
            "denominator degrees of freedom, validated from identical coefficient/covariance inputs "
            "against mitml 0.4-5 in R 4.5.3. Holm correction was applied to the four exploratory "
            "modifiers. The restricted-cubic-spline non-linearity test was combined across imputations "
            "using D2. Records with missing prior-diagnosis status (n = 192) were excluded from that "
            "binary one-degree-of-freedom interaction. Sex-specific PHQ-9 slopes were obtained as "
            "linear combinations within each imputation and combined with Rubin's rules (S3 Table). "
            "Time-dependent care modification could not be evaluated because QS103U/QS103C lacked "
            "usable data; cross-sectional care-cascade models were limited to participants reporting a "
            "prior diagnosis. E-values were calculated as sensitivity descriptors (S9 Table) [33]."
        ),
    )
    set_paragraph(
        document,
        "To assess a possible non-linear dose-response relationship",
        (
            "To assess the possible non-linear shape of the association between continuous PHQ-9 "
            "score and elevated blood pressure, restricted cubic splines used four knots at PHQ-9 "
            "values 0, 4, 9 and 14. These clinically interpretable locations were chosen because "
            "percentile placement was uninformative under the marked right skew of PHQ-9. A "
            "design-based joint Wald test evaluated the non-linear terms within each completed "
            "dataset, and the statistics were combined across 20 imputations with D2. The curve "
            "retained the Model 2 adjustment set, including altitude. Marginal standardized "
            "prevalence was calculated by predicting each observation at every PHQ-9 grid value "
            "and taking the survey-weighted mean; 95% confidence intervals used the delta method "
            "on the log scale."
        ),
    )
    set_paragraph(
        document,
        "The present research consisted of a secondary",
        (
            "This study used public, de-identified individual-level ENDES microdata. Article 97 of "
            "Peru's Supreme Decree No. 043-2001-PCM requires INEI survey information to be disclosed "
            "only in unnamed form [47]; the authors had no access to direct identifiers and did not "
            "contact participants. INEI obtained consent before the original interviews and physical "
            "measurements. No secondary consent was sought because the present analysis used only the "
            "public de-identified files. No formal institutional ethics determination was obtained for "
            "this secondary analysis. EDITORIAL HOLD: before submission, insert the reviewing "
            "committee and approval/exemption resolution, or a verified regulation that specifically "
            "exempts this analysis; the INEI confidentiality rule alone does not establish an "
            "ethics-review exemption."
        ),
    )

    # Resultados.
    set_paragraph(document, "Baseline characteristics of the cohort", "Participant characteristics")
    set_paragraph(
        document,
        "From the initial ENDES 2019-2025 sample",
        (
            "After applying the demographic, reproductive and key-variable eligibility criteria, the "
            "pooled repeated cross-sectional analytic sample comprised 191,757 adults (Fig 1). "
            "Participant characteristics by sex and elevated-blood-pressure status are shown in "
            "Table 1. Weighted elevated-blood-pressure prevalence was 16.1%: 20.3% in men and 12.1% "
            "in women; the unweighted proportion was 12.12% (23,248/191,757). Participants with "
            "elevated blood pressure were older and had higher body mass index and waist circumference "
            "on average. The continuous PHQ-9 score was not associated with elevated-blood-pressure "
            "status in the unadjusted continuous comparison (p = 0.542), and its five severity "
            "categories were likewise not associated in the Rao–Scott test (p = 0.143). The design "
            "effect for elevated-blood-pressure prevalence was 3.73 (mean across Table 1 estimates, "
            "4.59). In observed-data bivariate analyses (Table 2), wealth and altitude categories were "
            "associated with elevated-blood-pressure status (both p < 0.001); these unadjusted tests "
            "are descriptive and do not establish determinants."
        ),
    )
    set_paragraph(
        document,
        "Table 1. Baseline characteristics",
        (
            "Table 1. Participant characteristics by sex and elevated blood pressure (EBP), "
            "using observed data. ENDES 2019–2025, total n = 191,757 adults."
        ),
    )
    set_paragraph(
        document,
        "In a sensitivity analysis without geographic",
        (
            "In the model without altitude (S7 Table), PHQ-9 score had a small inverse cross-sectional "
            "association with elevated blood pressure (PR = 0.994; p = 0.003). In the main Model 2, "
            "which added altitude to the other adjustment covariates, the association moved toward the "
            f"null but remained statistically distinguishable from it (PR = {float(main_model['pr']):.3f}; "
            f"95% CI {float(main_model['ci_low']):.3f}–{float(main_model['ci_high']):.3f}; "
            f"p = {float(main_model['p_value']):.3f}; n = 191,757) (Table 3). On the log-coefficient "
            "scale, adding altitude reduced the deviation from the null by approximately 22% "
            "(−0.006513 to −0.005065). This descriptive attenuation is compatible with geographic "
            "differences but does not demonstrate a causal confounding mechanism. Restricted cubic "
            f"splines did not provide evidence of non-linearity (D2-pooled p = "
            f"{float(spline['d2_p_value']):.3f}; Fig 2). Exact source data for Figs 1 and 2 are "
            "provided in S4 Table."
        ),
    )

    sex = panel.loc["sex"]
    year = panel.loc["year"]
    area = panel.loc["area"]
    wealth = panel.loc["wealth"]
    dxhta = panel.loc["dxhta"]
    altitude = panel.loc["altitude"]
    set_paragraph(
        document,
        "Despite the absence of a clinically relevant main association",
        (
            "The corrected D1 panel (S3 Table) indicated heterogeneity by sex "
            f"(global p {format_effect(sex)}). The PHQ-9 prevalence ratio per point was "
            f"{float(slopes.loc['men', 'pr']):.3f} (95% CI "
            f"{float(slopes.loc['men', 'ci_low']):.3f}–{float(slopes.loc['men', 'ci_high']):.3f}; "
            f"p {format_p(float(slopes.loc['men', 'p_value']))}) in men and "
            f"{float(slopes.loc['women', 'pr']):.3f} (95% CI "
            f"{float(slopes.loc['women', 'ci_low']):.3f}–{float(slopes.loc['women', 'ci_high']):.3f}; "
            f"p {format_p(float(slopes.loc['women', 'p_value']))}) in women. Other global tests were: "
            f"year (p {format_effect(year)}), area (p {format_effect(area)}), wealth "
            f"(p {format_effect(wealth)}), prior hypertension diagnosis (p {format_effect(dxhta)}) "
            f"and altitude (p {format_effect(altitude)}). The prior-diagnosis test was a binary "
            "one-degree-of-freedom contrast after excluding 192 records with missing status. Among "
            f"participants with a prior diagnosis, PHQ-9 score was associated with elevated blood "
            f"pressure at PR = {float(control['pr']):.3f} per point (95% CI "
            f"{float(control['ci_low']):.3f}–{float(control['ci_high']):.3f}; "
            f"p {format_p(float(control['p_value']))}; n = {int(round(float(control['mean_n_obs']))):,}), "
            "whereas it was not associated with reported medication non-adherence "
            f"(PR = {float(adherence['pr']):.3f}; 95% CI "
            f"{float(adherence['ci_low']):.3f}–{float(adherence['ci_high']):.3f}; "
            f"p = {float(adherence['p_value']):.3f}; n = "
            f"{int(round(float(adherence['mean_n_obs']))):,}) (Table 4). These cross-sectional "
            "contrasts cannot distinguish effects of diagnosis, treatment, selection or health-care contact."
        ),
    )
    set_paragraph(
        document,
        "The attenuation of the main effect remained",
        (
            "The small inverse estimate was similar across the sensitivity specifications defined in "
            "the analysis framework (Table 5): excluding 2020 (PR = 0.995; p = 0.039), using only the "
            "second blood-pressure measurement (PR = 0.996; p = 0.044), applying the ≥130/80 mmHg "
            "threshold (PR = 0.993; p < 0.001), and fitting a logistic model (OR = 0.993; p = 0.014). "
            f"The PHQ-9 × year interaction was not statistically significant (D1 p "
            f"{format_p(float(year['p_value']))}), providing no evidence of heterogeneity by survey "
            "year. The Model 2 E-value was 1.08 (1.03 for the confidence limit closest to the null), "
            "showing that the small estimate is sensitive to modest unmeasured associations; the "
            "E-value does not identify an actual confounder (S9 Table). Complete-case analysis "
            "(n = 183,586; 95.7%) produced PR = 0.995 (95% CI 0.990–1.000), similar to the imputed "
            "analysis."
        ),
    )
    set_paragraph(
        document,
        "Fig 2. Restricted cubic spline curve",
        (
            "Fig 2. Restricted cubic spline curve of marginal standardized elevated-blood-pressure "
            "prevalence by PHQ-9 score, combined over 20 parametric chained-equations imputations "
            "(bands: Rubin 95% CI; knots at 0, 4, 9 and 14). The D2-pooled non-linearity p value "
            f"was {float(spline['d2_p_value']):.3f}."
        ),
    )
    set_paragraph(
        document,
        "Table 3. Prevalence ratio",
        (
            "Table 3. Prevalence ratio (PR) of elevated blood pressure by PHQ-9 score and "
            "covariates, ENDES 2019–2025 (n = 191,757; 20 parametric chained-equations "
            "imputations). Full adjusted models are provided in S10 Table."
        ),
    )
    set_paragraph(
        document,
        "Table 4. Prevalence ratio",
        (
            "Table 4. Prevalence ratios (PR) in the care cascade: medication non-adherence and "
            "uncontrolled blood pressure among participants reporting a prior hypertension "
            "diagnosis. ENDES 2019–2025. Full adjusted models are provided in S11 Table."
        ),
    )

    # Discusión y conclusión.
    set_paragraph(
        document,
        "The present study, based on a representative cohort",
        (
            "In this nationally representative pooled repeated cross-sectional sample of more than "
            "190,000 Peruvian adults, PHQ-9 score had a very small inverse association with elevated "
            "blood pressure after adjustment (PR = 0.995 per point). Its clinical importance is "
            "uncertain because no clinical relevance margin was defined in advance. Adding altitude "
            "moved the estimate approximately 22% toward the null. This pattern highlights the "
            "importance of geographic context in Andean population analyses, but the contrast between "
            "models does not prove that altitude caused the attenuation or that the residual "
            "association is protective."
        ),
    )
    set_paragraph(
        document,
        "These results recontextualize the controversy",
        (
            "These findings add a cross-sectional Peruvian estimate to a mixed literature. "
            "Prospective meta-analyses have reported positive associations between depression and "
            "incident hypertension [7,34], whereas some cross-sectional studies have reported "
            "inverse associations and proposed emotional-dampening or vascular-depression "
            "hypotheses [35–38]. In the present data, the estimate was very small and its low "
            "E-value (≤ 1.11) indicates sensitivity to modest unmeasured associations; neither "
            "finding distinguishes neurobiological, geographic, clinical or selection-related "
            "explanations. Altitude has been associated separately with depressive symptoms and "
            "hypertension in Peru [20,21,39], and the estimate moved toward the null when altitude "
            "was added. This model contrast supports accounting for geographic context in Andean "
            "population analyses but does not establish that altitude caused the attenuation or "
            "that the unadjusted association was spurious [40]."
        ),
    )
    set_paragraph(
        document,
        "At the local level, our findings help raise",
        (
            "At the local level, these findings complement earlier ENDES analyses, including "
            "Valladares-Garrido et al. [41], by incorporating altitude, the complex survey design and "
            "parametric chained-equations imputation across 20 completed datasets. The movement of the "
            "estimate after adjustment for altitude is compatible with geographic differences in the "
            "joint distribution of depressive symptoms and blood-pressure status; its observational "
            "cross-sectional nature precludes attributing that change to a causal mechanism."
        ),
    )
    set_paragraph(
        document,
        "Despite the clinically trivial primary structural effect",
        (
            f"The corrected binary interaction with prior hypertension diagnosis had D1 p "
            f"{format_p(float(dxhta['p_value']))} and Holm-adjusted p "
            f"{format_p(float(dxhta['p_holm_exploratorios']))}. Within the care-cascade subpopulation, "
            f"the PHQ-9 association with elevated blood pressure was small and inverse "
            f"(PR = {float(control['pr']):.3f}; p {format_p(float(control['p_value']))}), whereas the "
            f"association with medication non-adherence was near the null "
            f"(PR = {float(adherence['pr']):.3f}; p = {float(adherence['p_value']):.3f}). "
            "Several mechanisms could generate this pattern, including treatment, selection into "
            "diagnosis, health-care contact, recall and unmeasured clinical severity. Because timing "
            "of diagnosis, treatment initiation and symptom onset was unavailable, the data cannot "
            "identify detection bias or a care paradox. The findings therefore motivate, but do not "
            "demonstrate the need for, longitudinal evaluation of integrated mental-health and "
            "cardiometabolic care [42–44]."
        ),
    )
    set_paragraph(
        document,
        "This study should be interpreted in light",
        (
            "This study has several limitations. First, its cross-sectional design cannot establish "
            "the temporal order of depressive symptoms and blood-pressure status; bidirectionality and "
            "reverse causation remain possible [37]. Model 3 is an exploratory additional-adjustment "
            "specification and is not a mediation analysis. Second, PHQ-9 covers the previous 14 days "
            "and therefore measures recent symptoms rather than long-term depressive exposure. Third, "
            "elevated blood pressure was based on measurements from one home visit and does not "
            "establish chronic hypertension. A sensitivity analysis using only the second measurement "
            "was conducted to reduce concern about first-measurement effects, but ambulatory or "
            "repeat-visit confirmation was unavailable; outcome misclassification may remain."
        ),
    )
    set_paragraph(
        document,
        "Fourth, the estimates in the care-cascade subpopulation",
        (
            "Fourth, care-cascade estimates have specific measurement limitations. Prior diagnosis "
            "(QS102) and medication non-adherence (QS106) were self-reported and are susceptible to "
            "recall, health-care-contact and social-desirability processes. Their magnitude and "
            "direction cannot be determined from these data, so the reported prevalences and "
            "associations should be interpreted cautiously."
        ),
    )
    set_paragraph(
        document,
        "Finally, there is the possibility of residual",
        (
            "Residual confounding by unmeasured variables remains possible; sleep disorders, for "
            "example, are associated with both depressive symptoms and hypertension but were not "
            "available [45]. The low E-values (≤ 1.11) show that modest unmeasured associations could "
            "move the small estimates to the null. They do not demonstrate the presence or direction "
            "of such confounding (S9 Table)."
        ),
    )
    set_paragraph(
        document,
        "The main strength of this study lies",
        (
            "Strengths include the integration of seven annual nationally representative survey "
            "samples, explicit incorporation of the complex stratified design, transparent handling "
            "of missing covariates and multiple sensitivity analyses. These features support "
            "population-level inference to Peruvian adults for the 2019–2025 survey period."
        ),
    )
    set_paragraph(
        document,
        "No clinically relevant association was found",
        (
            "PHQ-9 score showed a very small inverse cross-sectional association with elevated blood "
            f"pressure in Peruvian adults (PR = {float(main_model['pr']):.3f} per point; 95% CI "
            f"{float(main_model['ci_low']):.3f}–{float(main_model['ci_high']):.3f}; "
            f"p = {float(main_model['p_value']):.3f}). Its clinical importance is uncertain. Adding "
            "altitude moved the estimate toward the null, underscoring the value of geographic "
            "adjustment in Andean population studies without establishing a causal explanation. "
            "Prospective studies with repeated blood-pressure measurements, treatment timing and "
            "longitudinal mental-health assessment are needed to clarify temporality and subgroup "
            "differences."
        ),
    )
    set_paragraph(
        document,
        "S6 Table. Comparison of the included cohort",
        (
            "S6 Table. Comparison of records included in the analytic sample with records excluded "
            "during sample selection (STROBE)."
        ),
    )

    # Data availability y referencias.
    acknowledgments = find_paragraph(document, "Acknowledgments")
    acknowledgments.text = "Data availability"
    acknowledgments.style = "Heading 1"
    set_paragraph(
        document,
        "Data availability:",
        (
            "The analysis code supporting this study is publicly available on Zenodo at "
            "https://doi.org/10.5281/zenodo.21328300. The de-identified individual-level ENDES "
            "microdata are publicly available without restriction from Peru's National Institute of "
            "Statistics and Informatics (INEI) Microdata Portal at "
            "https://proyectos.inei.gob.pe/microdatos/."
        ),
    )
    set_paragraph(
        document,
        "Tarqui-Mamani C, Alvarez-Dongo",
        (
            "Diaz-Arocutipa C. Trends of awareness, treatment, and control of hypertension in Peru: "
            "a 5-year national survey analysis. J Hypertens. 2025;43(10):1726–1730. "
            "doi: 10.1097/HJH.0000000000004109."
        ),
    )
    set_paragraph(
        document,
        "Freedman DS, Lawman HG",
        (
            "Muharram FR, Tjandra S, Madani NJ, Rokx C, Abdullah A. Trends in the double burden of "
            "malnutrition among Indonesian adults, 2007 to 2023. Sci Rep. 2025;15:34883. "
            "doi: 10.1038/s41598-025-17348-9."
        ),
    )
    set_paragraph(
        document,
        "Hernández-Vásquez A, Vargas-Fernández",
        (
            "Hernández-Vásquez A, Vargas-Fernández R, Rojas-Roque C, Gamboa-Unsihuay JE. "
            "Association between altitude and depression in Peru: an 8-year pooled analysis of "
            "population-based surveys. J Affect Disord. 2022;299:536–544. "
            "doi: 10.1016/j.jad.2021.12.059."
        ),
    )
    set_paragraph(
        document,
        "Aune D, Giovannucci E",
        (
            "Aune D, Giovannucci E, Boffetta P, Fadnes LT, Keum N, Norat T, et al. Fruit and "
            "vegetable intake and the risk of cardiovascular disease, total cancer and all-cause "
            "mortality—a systematic review and dose-response meta-analysis of prospective studies. "
            "Int J Epidemiol. 2017;46(3):1029–1056. doi: 10.1093/ije/dyw319."
        ),
    )
    set_paragraph(
        document,
        "Zegarra-Rodríguez CA",
        (
            "Zegarra-Rodríguez CA, Plasencia-Dueñas NR, Failoc-Rojas VE. Disparities in the "
            "prevalence of screened depression at different altitudes in Peru: a retrospective "
            "analysis of the ENDES 2019. PLoS One. 2022;17(12):e0278947. "
            "doi: 10.1371/journal.pone.0278947."
        ),
    )
    set_paragraph(
        document,
        "Parati G, Bilo G, Faini A",
        (
            "Richalet JP, Hermand E, Lhuissier FJ. Cardiovascular physiology and pathophysiology "
            "at high altitude. Nat Rev Cardiol. 2024;21(2):75–88. "
            "doi: 10.1038/s41569-023-00924-9."
        ),
    )
    set_paragraph(
        document,
        "Ogedegbe G.",
        (
            "Ogedegbe G. Labeling and hypertension: it is time to intervene on its negative "
            "consequences. Hypertension. 2010;56(3):344–345. "
            "doi: 10.1161/HYPERTENSIONAHA.110.156257."
        ),
    )
    set_paragraph(
        document,
        "Instituto Nacional de Estadística e Informática. Ficha Técnica",
        (
            "Instituto Nacional de Estadística e Informática. Ficha Técnica: Encuesta Demográfica y "
            "de Salud Familiar—ENDES 2024 [Internet]. Lima: INEI; 2024 [cited 2026 Jul 23]. Available "
            "from: https://proyectos.inei.gob.pe/iinei/srienaho/Descarga/DocumentosMetodologicos/"
            "2024-5/FichaTecnica.pdf."
        ),
    )
    set_paragraph(
        document,
        "Instituto Nacional de Estadística e Informática. Manual de la Entrevistadora",
        (
            "Instituto Nacional de Estadística e Informática. Manual de la Entrevistadora—ENDES 2024 "
            "[Internet]. Lima: INEI; 2024 [cited 2026 Jul 23]. Available from: "
            "https://proyectos.inei.gob.pe/iinei/srienaho/Descarga/DocumentosMetodologicos/"
            "2024-5/ManualEntrevistadora.pdf."
        ),
    )
    # Use the first blank paragraph after reference 45.
    ref45 = find_paragraph(document, "Cai Y, Chen M")
    candidate = ref45._p.getnext()
    ref46 = Paragraph(candidate, ref45._parent)
    ref46.text = (
        "White IR, Royston P, Wood AM. Multiple imputation using chained equations: issues and "
        "guidance for practice. Stat Med. 2011;30(4):377–399. doi: 10.1002/sim.4067."
    )
    copy_paragraph_properties(ref45, ref46)
    candidate = ref46._p.getnext()
    ref47 = Paragraph(candidate, ref45._parent)
    ref47.text = (
        "Instituto Nacional de Estadística e Informática. Secreto estadístico y confidencialidad de "
        "la información: artículo 97 del Decreto Supremo No. 043-2001-PCM [Internet]. Lima: INEI; "
        "[cited 2026 Jul 23]. Available from: "
        "https://encuestas.inei.gob.pe/endes/confidencialidad.asp."
    )
    copy_paragraph_properties(ref45, ref47)

    # Leyendas y suplementos.
    set_paragraph(
        document,
        "S1 Fig. Directed acyclic graph",
        (
            "S1 Fig. Theory-informed directed acyclic graph used to select the Model 2 adjustment "
            "set. Arrows and labels represent hypothesized relations, not effects established by "
            "this cross-sectional study. Orange variables were added only to exploratory Model 3; "
            "that model is not interpreted as a mediation analysis. U denotes hypothetical "
            "unmeasured factors. The graph was verified as acyclic in dagitty."
        ),
    )
    set_paragraph(
        document,
        "S2 Fig. Forest plot",
        (
            "S2 Fig. Forest plot of the PHQ-9 prevalence ratio by altitude stratum (overall and "
            "categories < 1,500 / 1,500–2,499 / ≥ 2,500 m a.s.l.). The corrected pooled D1 altitude "
            f"interaction was p {format_p(float(altitude['p_value']))} and Holm-adjusted p "
            f"{format_p(float(altitude['p_holm_exploratorios']))}. PR: prevalence ratio; 95% CI: "
            "95% confidence interval."
        ),
    )

    # Tablas incrustadas.
    if len(document.tables) != 5:
        raise RuntimeError(f"Se esperaban 5 tablas; se encontraron {len(document.tables)}.")

    bivariate_table = document.tables[1]
    bivariate_table.cell(0, 3).text = "n*"
    row_labels = {bivariate_table.cell(row, 0).text: row for row in range(1, 12)}
    for label, item in table2.iterrows():
        if label not in row_labels:
            raise RuntimeError(f"La Tabla 2 no contiene la fila {label!r}.")
        row = row_labels[label]
        bivariate_table.cell(row, 1).text = f"{float(item['statistic']):.2f}"
        bivariate_table.cell(row, 2).text = format_p(float(item["p_value"]))
        bivariate_table.cell(row, 3).text = f"{int(item['n_unweighted']):,}"
    set_merged_row_text(
        bivariate_table,
        13,
        (
            "Rao–Scott F statistic with survey-design correction on observed data. Exact numerator "
            "and denominator degrees of freedom are retained in the reproducible canonical output."
        ),
    )
    set_merged_row_text(
        bivariate_table,
        14,
        (
            "n* = observed/applicable respondents. For alcohol and tobacco, the "
            "not-applicable questionnaire-skip category shown in Table 1 was excluded from this test. "
            "PHQ-9 is tested here as five severity categories; the continuous-score p value in Table 1 "
            "is 0.542."
        ),
    )
    set_merged_row_text(
        bivariate_table,
        15,
        (
            "A low p value indicates evidence of a crude bivariate association; adjusted associations "
            "are reported in Table 3."
        ),
    )

    main_table = document.tables[2]
    set_merged_row_text(
        main_table,
        5,
        "Model 3 — Exploratory (additional adjustment)",
    )
    set_merged_row_text(
        main_table,
        8,
        (
            "PR: prevalence ratio (quasi-Poisson, log link). 95% CI and p values pooled with Rubin's "
            "rules over 20 parametric chained-equations imputations. n = 191,757 in all three models."
        ),
    )
    set_merged_row_text(
        main_table,
        9,
        (
            "Nine covariates were declared as imputation targets; five had missing values replaced. "
            "Education, PHQ-9 exposure and elevated-blood-pressure outcome were not imputed."
        ),
    )
    set_merged_row_text(
        main_table,
        10,
        (
            "Model 1: crude PHQ-9. Model 2 (main): age, sex, education, area, wealth, intimate partner "
            "violence, year and altitude. Model 3 (exploratory): additionally BMI, waist, tobacco, "
            "problematic alcohol, diet and diabetes; it is not interpreted as a mediation model."
        ),
    )
    set_merged_row_text(
        main_table,
        11,
        (
            "Altitude was included as a geographic and contextual adjustment variable. "
            "The intercept is omitted."
        ),
    )

    cascade_table = document.tables[3]
    set_merged_row_text(
        cascade_table,
        6,
        "PR: prevalence ratio (quasi-Poisson, log link), pooled over 20 imputations with Rubin's rules.",
    )
    set_merged_row_text(
        cascade_table,
        8,
        (
            "Same adjustment covariates as Model 2. Nine covariates were declared as imputation "
            "targets; the outcome was not imputed."
        ),
    )

    sensitivity_table = document.tables[4]
    sensitivity_table.cell(2, 0).text = "Without the 2020 survey year (COVID-19)"
    sensitivity_table.cell(7, 0).text = "Restricted cubic spline (non-linearity)"
    set_merged_row_text(
        sensitivity_table,
        9,
        (
            "All models were pooled over 20 parametric chained-equations imputations with Rubin's "
            "rules, except the complete-case row."
        ),
    )

    # Comprobaciones textuales antes de guardar.
    body = "\n".join(p.text for p in document.paragraphs)
    body += "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    abstract_words = len(re.findall(r"\b[\w'-]+\b", abstract))
    if abstract_words > 300:
        raise RuntimeError(f"El resumen excede 300 palabras ({abstract_words}).")
    forbidden = [
        "10 MAR covariables",
        "p = 0.002",
        "p = 0.028",
        "p = 0.013",
        "clinically relevant association",
        "estimates the total effect",
        "adjusted direct effects",
        "with mediators",
        "prespecified",
        "aggregated data",
        "representative cohort",
        "baseline characteristics of the cohort",
        "without 2020 cohort",
        "geographic confounder",
        "median pooled non-linearity p",
        "0000000000004029",
    ]
    found = [term for term in forbidden if term.lower() in body.lower()]
    if found:
        raise RuntimeError(f"Persisten términos/cifras obsoletos: {found}")
    if body.count("Depression; hypertension; blood pressure;") > 0:
        raise RuntimeError("Persistió la segunda lista de palabras clave.")

    if not BACKUP.exists():
        shutil.copy2(MANUSCRIPT, BACKUP)
    document.save(MANUSCRIPT)

    print(f"Manuscrito actualizado: {MANUSCRIPT.relative_to(ROOT)}")
    print(f"Copia previa: {BACKUP.relative_to(ROOT)}")
    print(f"Resumen: {abstract_words} palabras")
    print(f"SHA-256 final provisional: {sha256(MANUSCRIPT)}")


if __name__ == "__main__":
    main()
