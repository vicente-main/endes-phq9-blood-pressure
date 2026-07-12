# -*- coding: utf-8 -*-
"""Aplica los pendientes 1-7 de PENDIENTES_MANUSCRITO.md al manuscrito .docx.

Cada edicion verifica que el texto objetivo exista (assert) antes de tocarlo, de
modo que el script falla ruidosamente si el documento cambio de estructura.
"""
from __future__ import annotations

import copy
from pathlib import Path

import docx
import docx.oxml.ns
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[2]
DOC = ROOT / "SEVERIDAD SINTOMAS DEPRESIVOS VS PRESION ARTEIAL ELEVADA (1).docx"


def replace_in_run(run, old, new, *, required=True):
    if old in run.text:
        run.text = run.text.replace(old, new)
        return True
    if required:
        raise AssertionError(f"No se encontro el texto a reemplazar: {old!r}\n  en: {run.text!r}")
    return False


def append_run_like(paragraph, neighbor_run, text):
    """Agrega un run al final del parrafo copiando el formato (rPr) de neighbor_run."""
    new_run = paragraph.add_run(text)
    if neighbor_run is not None and neighbor_run._element.rPr is not None:
        new_run._element.append(copy.deepcopy(neighbor_run._element.rPr))
    return new_run


def add_reference_after(last_ref_para, text):
    """Clona el ultimo item (autonumerado) de la lista de referencias y le pone texto nuevo."""
    new_p = copy.deepcopy(last_ref_para._p)
    # Limpia todos los runs del clon.
    for r in list(new_p.findall(qn("w:r"))):
        new_p.remove(r)
    last_ref_para._p.addnext(new_p)
    para = Paragraph(new_p, last_ref_para._parent)
    # Reusa el formato del primer run original si existe.
    src_run = last_ref_para.runs[0] if last_ref_para.runs else None
    append_run_like(para, src_run, text)
    return para


def main():
    d = docx.Document(str(DOC))
    P = d.paragraphs

    # ====== ITEM 1: definicion del desenlace + consistencia hemodinamica (PAS>PAD) ======
    p52 = P[52]
    assert p52.runs[1].text.strip() == "(1).", p52.runs[1].text
    item1 = (
        " La presion arterial se determino a partir de dos tomas realizadas en una unica "
        "visita domiciliaria (variables QS903 y QS905); se definio presion arterial elevada "
        "como un promedio de presion arterial sistolica >= 140 mmHg o de presion arterial "
        "diastolica >= 90 mmHg. Antes de promediar las tomas se depuraron las mediciones: se "
        "descartaron los valores centinela (999) y los fisiologicamente implausibles (sistolica "
        "< 70 o > 270 mmHg; diastolica < 30 o > 150 mmHg), y se excluyeron por inconsistencia "
        "hemodinamica los registros con inversion tensional -aquellos en que la presion sistolica "
        "no superaba a la diastolica en alguna de las tomas-, que por tanto no aportaron un "
        "desenlace valido."
    )
    append_run_like(p52, p52.runs[1], item1)

    # ====== ITEM 4: citas de altitud a estilo numerico (26, 27) ======
    p54 = P[54]
    assert "Hern" in p54.runs[4].text and "Mendoza-Quispe" in p54.runs[4].text, p54.runs[4].text
    p54.runs[4].text = " (26, 27)"

    # ====== ITEM 3: remitir confusores (no solo altitud) a Figura S2 / Tabla S5 ======
    item3 = (
        " La justificacion causal y el respaldo bibliografico de cada confusor estructural "
        "-y de las aristas confusor->exposicion y confusor->desenlace que motivan su ajuste- "
        "se documentan arista por arista en el grafo aciclico dirigido (Figura S2) y en la Tabla S5."
    )
    append_run_like(p54, p54.runs[-1], item3)

    # ====== ITEM 2: citas Lumley (survey) y Rubin (reglas de Rubin) en Analisis estadistico ======
    p59 = P[59]
    replace_in_run(p59.runs[0], "el paquete survey de R", "el paquete survey (28) de R")
    assert p59.runs[6].text.rstrip().endswith("reglas de Rubin"), p59.runs[6].text
    p59.runs[6].text = p59.runs[6].text + " (29)"

    # ====== ITEM 6a: prevalencia ponderada correcta (12,13% no ponderada -> 16,2% ponderada) ======
    p68 = P[68]
    r = p68.runs[0]
    replace_in_run(
        r,
        "La prevalencia ponderada de presión arterial elevada fue de 12,13%, siendo significativamente mayor en hombres (16,6%) que en mujeres (8,7%).",
        "La prevalencia ponderada de presión arterial elevada fue de 16,2 %, significativamente mayor en hombres (20,4 %) que en mujeres (12,2 %); la proporción no ponderada en la muestra analítica fue de 12,13 % (19 985/164 719).",
    )
    # ====== ITEM 6b: DEFF (3,91 -> 3,66) + reparar fragmento "ajuste d...e varianza empleado" ======
    replace_in_run(
        r,
        "El efecto de diseño (DEFF) global de las estimaciones fue de 3,91, lo que justifica la necesidad del ajuste dPrevalencia bivariada El análisis bivariado",
        "El efecto de diseño (DEFF) de la prevalencia de presión arterial elevada fue de 3,66 (DEFF promedio de las estimaciones de la Tabla 1 = 4,55), lo que justifica la necesidad del ajuste de varianza empleado. Prevalencia bivariada. El análisis bivariado",
    )

    # ====== ITEM 6c: cascada de cuidado en Resultados (claim de no adherencia es incorrecto) ======
    p72 = P[72]
    r72 = p72.runs[0]
    replace_in_run(
        r72,
        "Al explorar este fenómeno en el sub-modelo de cascada de cuidado (Tabla 4), restringido a los 14 956 individuos previamente diagnosticados, la mayor severidad de síntomas depresivos se asoció a un incremento en el riesgo de no adherencia al tratamiento antihipertensivo farmacológico, demostrando el impacto conductual negativo de la comorbilidad psiquiátrica en pacientes ya detectados por el sistema de salud.e varianza empleado.",
        "Al explorar este fenómeno en los sub-modelos de cascada de cuidado (Tabla 4), restringidos a los individuos con diagnóstico médico previo de hipertensión, la mayor severidad de síntomas depresivos se asoció a una menor prevalencia de presión arterial elevada -es decir, mejor control tensional aparente- en los 14 956 pacientes diagnosticados (RP = 0,99 por punto del PHQ-9; IC 95 % 0,98-1,00; p = 0,002). En cambio, la severidad depresiva no se asoció de forma significativa con la no adherencia al tratamiento antihipertensivo (RP = 1,00; IC 95 % 0,99-1,01; p = 0,97) en los 14 943 pacientes con datos de adherencia. Este patrón es coherente con un sesgo de detección/cuidado más que con un deterioro conductual de la adherencia.",
    )

    # ====== ITEM 6d: cascada en Discusion (mismo claim incorrecto) ======
    p79 = P[79]
    r79 = p79.runs[0]
    replace_in_run(
        r79,
        "En la subpoblación de la cascada de cuidado, la severidad depresiva actuó como un potente predictor de la no adherencia al tratamiento farmacológico, hallazgo congruente con la literatura que documenta el impacto negativo de la psiquiatría en el autocuidado.",
        "En la subpoblación de la cascada de cuidado, los pacientes con diagnóstico previo y mayor severidad depresiva mostraron una prevalencia algo menor de presión arterial elevada (RP = 0,99; p = 0,002), mientras que la severidad depresiva no predijo de manera significativa la no adherencia al tratamiento (RP = 1,00; p = 0,97). Este patrón es más compatible con un sesgo de detección/cuidado -los pacientes detectados y tratados tienden a presentar una presión arterial mejor controlada y, a la vez, mayor carga psicológica- que con un deterioro conductual de la adherencia.",
    )

    # ====== ITEM 7: E-values en Metodos, Resultados y Limitaciones ======
    p63 = P[63]
    item7_met = (
        " Finalmente, para cuantificar la sensibilidad de los estimadores a la confusion no "
        "medida se calcularon E-values (VanderWeele y Ding, 2017) (30) para la razon de prevalencia "
        "del PHQ-9 en los modelos principales y en la cascada de cuidado (Tabla S9)."
    )
    append_run_like(p63, p63.runs[-1], item7_met)

    p74 = P[74]
    item7_res = (
        " El analisis de E-values reforzo esta interpretacion: dado que el intervalo de "
        "confianza del Modelo 2 incluye la unidad, no se requiere confusion no medida alguna para "
        "ser compatible con la ausencia de efecto, y la pequena asociacion inversa observada sin "
        "ajuste por altitud tendria un E-value de apenas 1,08 (1,04 para el limite del intervalo "
        "de confianza), de modo que un confusor tan debil como la propia altitud basta para "
        "explicarla por completo (Tabla S9)."
    )
    append_run_like(p74, p74.runs[-1], item7_res)

    p81 = P[81]
    item7_lim = (
        " No obstante, la magnitud de esta confusion potencial es acotada: los E-values de los "
        "estimadores principales fueron muy bajos (<= 1,11), lo que indica que asociaciones no "
        "medidas modestas bastarian para explicar las pequenas senales observadas, en linea con "
        "la interpretacion de que no existe un efecto causal independiente (Tabla S9)."
    )
    append_run_like(p81, p81.runs[-1], item7_lim)

    # ====== ITEMS 2/4/7: agregar las 5 referencias nuevas (26-30) al final de la lista ======
    last_ref = P[108]
    assert last_ref.text.startswith("Aune D"), last_ref.text
    refs = [
        "Hernández-Vásquez A, Vargas-Fernández R, Rojas-Roque C, Gamboa-Unsihuay JE. Association between altitude and depression in Peru: an 8-year pooled analysis of population-based surveys. J Affect Disord. 2022;299:536-544. doi: 10.1016/j.jad.2021.12.039.",
        "Mendoza-Quispe D, Chambergo-Michilot D, Moscoso-Porras M, Bernabe-Ortiz A. Hypertension prevalence by degrees of urbanization and altitude in Peru: pooled analysis of 186 906 participants. J Hypertens. 2023;41(7):1142-1151. doi: 10.1097/HJH.0000000000003444.",
        "Lumley T. Complex Surveys: A Guide to Analysis Using R. Hoboken (NJ): John Wiley & Sons; 2010.",
        "Rubin DB. Multiple Imputation for Nonresponse in Surveys. New York: John Wiley & Sons; 1987.",
        "VanderWeele TJ, Ding P. Sensitivity analysis in observational research: introducing the E-value. Ann Intern Med. 2017;167(4):268-274. doi: 10.7326/M16-2607.",
    ]
    anchor = last_ref
    for txt in refs:
        anchor = add_reference_after(anchor, txt)

    # ====== ITEM 7 (suplemento): leyenda de Tabla S9 al final del documento ======
    cap = P[117]
    assert cap.runs[0].text.strip().startswith("Tabla S8"), cap.runs[0].text
    s9_para = d.add_paragraph()
    bold = s9_para.add_run("Tabla S9. ")
    bold.bold = True
    s9_para.add_run(
        "E-values de los estimadores principales (VanderWeele y Ding, 2017). Para cada razón de "
        "prevalencia (o de momios, en la sensibilidad logística) se reporta el E-value de la "
        "estimación puntual y el del límite del intervalo de confianza más cercano a la unidad. El "
        "E-value es la fuerza mínima de asociación (en escala de razón) que un confusor no medido "
        "debería tener con la exposición y con el desenlace, por encima de los confusores ya "
        "ajustados, para explicar por completo la asociación observada."
    )

    d.save(str(DOC))
    print("OK - manuscrito actualizado")
    # Reporte de verificacion: recargar y contar referencias.
    d2 = docx.Document(str(DOC))
    refs_n = sum(1 for p in d2.paragraphs if p._p.find(qn("w:pPr")) is not None
                 and p._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None)
    print("Items de lista (referencias) detectados:", refs_n)


if __name__ == "__main__":
    main()
